#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word样式模板生成器 - 精简版
创建自定义样式并确保正确排序
"""

import yaml
import os
import re
import time
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx2pdf import convert
import pythoncom
import win32com.client as win32
import comtypes.client

STYLE_MAPPING = {
    '正文': '\u200B正文',
    '图题': '\u200C图题',
    '一级标题': '\u200D一级标题',
    '二级标题': '\u200e二级标题',
    '三级标题': '\u200F三级标题',
    '参考文献': '\u206F参考文献',
    '全局标题': '\u2010全局标题',
    '全文大标题': '\u2010全文大标题',
    '参考文献标题': '\u2011参考文献标题'
}

def ensure_consistent_page_setup(doc, config=None):
    """确保文档中所有section都使用相同的A4页面设置"""
    # 获取页面设置配置
    page_setup = config.get('page_setup', {}) if config else {}
    
    for i, section in enumerate(doc.sections):
        section.page_height = Cm(page_setup.get('page_height', 29.7))  # A4高度
        section.page_width = Cm(page_setup.get('page_width', 21))     # A4宽度
        section.left_margin = Cm(page_setup.get('left_margin', 2.54))   # 左边距2.54cm
        section.right_margin = Cm(page_setup.get('right_margin', 2.54))  # 右边距2.54cm  
        section.top_margin = Cm(page_setup.get('top_margin', 2.54))    # 上边距2.54cm
        section.bottom_margin = Cm(page_setup.get('bottom_margin', 2.54)) # 下边距2.54cm
        section.gutter = Cm(page_setup.get('gutter', 0.5))        # 装订线0.5cm
        section.gutter_pos = page_setup.get('gutter_pos', 0)          # 装订线位置：靠左
        
        # 设置页眉页脚距离
        section.header_distance = Cm(page_setup.get('header_distance', 1.5))  # 页眉位置1.5cm
        section.footer_distance = Cm(page_setup.get('footer_distance', 1.75))  # 页脚位置1.75cm
        
        # 添加页眉内容
        if 'header_content' in page_setup:
            header = section.header
            header_para = header.paragraphs[0]
            # 清空现有内容
            header_para.clear()
            # 添加页眉文字
            header_run = header_para.add_run(page_setup['header_content'])
            # 设置页眉对齐方式
            header_para.alignment = get_alignment_constant(page_setup.get('header_alignment', ''))
            # 设置页眉字体
            header_font = page_setup.get('header_font', '')
            header_size = page_setup.get('header_font_size', 0)
            set_run_font(header_run, header_font, header_size)
            
            # 添加页眉边框：自定义样式 - 上面粗线，下面细线，总宽度3磅
            pPr = header_para._element.get_or_add_pPr()
            
            # 移除现有的边框设置
            existing_pBdr = pPr.find(qn('w:pBdr'))
            if existing_pBdr is not None:
                pPr.remove(existing_pBdr)
            
            # 创建新的边框设置
            pBdr = OxmlElement('w:pBdr')
            pPr.append(pBdr)
            
            # 自定义双线效果：使用Word内置的粗细双线样式
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'thinThickSmallGap')  # 粗细双线样式，上线粗下线细
            bottom.set(qn('w:sz'), '24')  # 总宽度3磅
            bottom.set(qn('w:color'), '000000')  # 黑色
            bottom.set(qn('w:space'), '0')  # 紧贴文字下方
            pBdr.append(bottom)

def convert_word_to_pdf(word_path, pdf_path):
    """将Word文档转换为PDF"""
    try:
        pythoncom.CoInitialize()
        word_app = win32.DispatchEx('Word.Application')
        word_app.Visible = False
        word_app.DisplayAlerts = False
        
        doc = word_app.Documents.Open(word_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word_app.Quit()
        pythoncom.CoUninitialize()
        return True
    except Exception as e:
        print(f"PDF转换失败: {e}")
        return False

# 全局变量：图片序号计数器
image_counter = 1
# 全局变量：表格序号计数器
table_counter = 1
# 全局变量：图片数量统计
small_image_count = 0
large_image_count = 0
# 全局变量：上下标数量统计
subscript_count = 0
superscript_count = 0

class TitleCollector:
    """标题收集器 - 用于收集文档中的所有标题信息"""
    
    def __init__(self):
        self.titles = []  # [(title_text, style_name, level, paragraph_index, page_number)]
        self.title_counter = {'chapter': 0, 'section': 0, 'subsection': 0}
        self.current_paragraph_index = 0
        self.page_breaks = []  # 记录分页符位置
    
    def reset(self):
        """重置收集器"""
        self.titles.clear()
        self.title_counter = {'chapter': 0, 'section': 0, 'subsection': 0}
        self.current_paragraph_index = 0
        self.page_breaks.clear()
    
    def add_title(self, title_text, style_name, paragraph_index, page_number=None):
        """添加标题记录"""
        if page_number is not None:
            # 如果提供了页码，直接使用
            estimated_page = page_number
        else:
            # 根据段落索引估算页码（假设每页约25个段落）
            estimated_page = self._estimate_page_from_paragraph(paragraph_index)
        
        # 确定标题级别
        level = self._determine_level(style_name)
        
        # 提取实际标题文本，去掉原始编号
        actual_title = self._extract_clean_title(title_text, level)
        
        # 从原始标题中提取编号，而不是生成新编号
        number_pattern = r'^(\d+(\.\d+)*)'
        number_match = re.match(number_pattern, title_text)
        if number_match:
            formatted_number = number_match.group(1)
        else:
            # 如果没有数字编号，使用中文编号
            if level == 1:
                self.title_counter['chapter'] += 1
                self.title_counter['section'] = 0  # 重置二级标题计数器
                self.title_counter['subsection'] = 0  # 重置三级标题计数器
                formatted_number = f"第{self._get_chinese_number(self.title_counter['chapter'])}章"
            elif level == 2:
                self.title_counter['section'] += 1
                self.title_counter['subsection'] = 0  # 重置三级标题计数器
                formatted_number = f"{self.title_counter['chapter']}.{self.title_counter['section']}"
            elif level == 3:
                self.title_counter['subsection'] += 1
                formatted_number = f"{self.title_counter['chapter']}.{self.title_counter['section']}.{self.title_counter['subsection']}"
            else:
                formatted_number = actual_title
        
        self.titles.append({
            'text': actual_title,
            'original_style': style_name,
            'level': level,
            'number': formatted_number,
            'paragraph_index': paragraph_index,
            'page_number': estimated_page
        })
        
        # 更新当前段落索引
        self.current_paragraph_index = max(self.current_paragraph_index, paragraph_index)
    
    def _extract_clean_title(self, title_text, level):
        """从标题文本中提取实际标题，去掉所有编号"""
        # 处理数字编号格式：1. X, 1.1 X, 1.1.1 X
        number_pattern = r'^(\d+(\.\d+)*)\s*(.+)$'
        match = re.match(number_pattern, title_text)
        if match:
            return match.group(3).strip()
        
        # 处理中文编号格式：第X章 X
        chapter_pattern = r'^第.+?章\s*(.+)$'
        match = re.match(chapter_pattern, title_text)
        if match:
            return match.group(1).strip()
        
        # 处理带空格的数字编号：1 . X
        spaced_number_pattern = r'^(\d+(\.\d+)*)\s+\.\s+(.+)$'
        match = re.match(spaced_number_pattern, title_text)
        if match:
            return match.group(3).strip()
        
        # 处理重复文本情况，如"图像布局示例 图像布局示例"
        # 检测连续重复的文本
        duplicate_pattern = r'^(.+)\s+\1$'
        match = re.match(duplicate_pattern, title_text)
        if match:
            return match.group(1).strip()
        
        # 如果没有匹配到编号格式，返回原文本
        return title_text
    

    
    def add_page_break(self, paragraph_index):
        """记录分页符位置"""
        self.page_breaks.append(paragraph_index)
    
    def _estimate_page_from_paragraph(self, paragraph_index):
        """根据段落索引估算页码"""
        # 统计在当前段落之前的分页符数量
        page_breaks_before = sum(1 for pb in self.page_breaks if pb <= paragraph_index)
        return page_breaks_before + 1
    
    def _determine_level(self, style_name):
        """根据样式名确定标题级别"""
        if style_name == '\u200D一级标题':  # 一级标题
            return 1
        elif style_name == '\u200e二级标题':  # 二级标题
            return 2
        elif style_name == '\u200F三级标题':  # 三级标题
            return 3
        else:
            return 0  # 未知级别
    
    def _get_chinese_number(self, num):
        """获取中文数字"""
        chinese_numbers = ['', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        if num <= 10:
            return chinese_numbers[num]
        elif num < 20:
            return '十' + chinese_numbers[num - 10]
        elif num % 10 == 0:
            return chinese_numbers[num // 10] + '十'
        else:
            return chinese_numbers[num // 10] + '十' + chinese_numbers[num % 10]
    
    def get_all_titles(self):
        """获取所有收集的标题"""
        return self.titles
    
    def get_titles_by_level(self, level):
        """按级别获取标题"""
        return [title for title in self.titles if title['level'] == level]

class TableOfContentsGenerator:
    """目录生成器"""
    
    def __init__(self, title_collector, toc_config):
        self.title_collector = title_collector
        self.toc_config = toc_config
        # 从配置文件加载样式
        self.toc_styles = {
            'toc_title': toc_config.get('toc_title', {}),
            'chapter_entry': toc_config.get('toc_entries', {}).get('chapter_entry', {}),
            'section_entry': toc_config.get('toc_entries', {}).get('section_entry', {}),
            'subsection_entry': toc_config.get('toc_entries', {}).get('subsection_entry', {})
        }
    
    def generate_toc(self, doc):
        """使用手动方式生成目录，确保格式正确"""
        # 添加目录标题
        toc_title_config = self.toc_styles['toc_title']
        toc_title = doc.add_paragraph()
        
        # 创建目录标题，确保字体均匀粗体
        toc_title_run = toc_title.add_run(toc_title_config.get('text', ''))
        set_run_font(toc_title_run, toc_title_config.get('font', ''), toc_title_config.get('size', 0), bold=toc_title_config.get('bold', False))
        
        # 设置段落格式
        toc_title.alignment = get_alignment_constant(toc_title_config.get('alignment', ''))
        toc_title.paragraph_format.space_before = Pt(toc_title_config.get('space_before', 0))
        toc_title.paragraph_format.space_after = Pt(toc_title_config.get('space_after', 0))
        toc_title.paragraph_format.line_spacing = toc_title_config.get('line_spacing', 0)
        
        # 重新计算所有标题的页码
        self._recalculate_page_numbers(doc)
        
        # 添加手动目录条目
        for title in self.title_collector.titles:
            self._add_toc_entry(doc, title)
        
        # 确保页面设置一致
        ensure_consistent_page_setup(doc)
        
        # 记录当前位置作为分页符
        self.title_collector.add_page_break(len(doc.paragraphs) - 1)
    
    def _recalculate_page_numbers(self, doc):
        """重新计算所有标题的页码"""
        # 简单的页码分配，根据标题级别和顺序分配页码
        # 从配置文件获取起始页码
        page_calc_config = self.toc_config.get('page_calculation', {})
        start_page = page_calc_config.get('start_page', 1)
        
        # 为了演示效果，我们根据标题顺序分配页码
        # 实际应用中可以根据段落位置和每页段落数进行更精确的计算
        for i, title in enumerate(self.title_collector.titles):
            # 简单分配：奇数标题从起始页开始，偶数标题加1
            # 或者根据标题级别调整页码
            if title['level'] == 1:
                # 一级标题从起始页开始
                title['page_number'] = start_page
            elif title['level'] == 2:
                # 二级标题与一级标题同页
                title['page_number'] = start_page
            else:
                # 三级标题与二级标题同页
                title['page_number'] = start_page
            
            # 每4个标题增加一页，模拟实际文档分页
            if (i + 1) % 4 == 0:
                start_page += 1
    
    def _add_toc_entry(self, doc, title):
        """添加单个目录条目"""
        para = doc.add_paragraph()
        
        # 根据级别选择样式
        if title['level'] == 1:
            style = self.toc_styles['chapter_entry']
        elif title['level'] == 2:
            style = self.toc_styles['section_entry']
        elif title['level'] == 3:
            style = self.toc_styles['subsection_entry']
        else:
            style = self.toc_styles['section_entry']
        
        # 设置段落格式
        para.alignment = get_alignment_constant(style.get('alignment', ''))
        para.paragraph_format.left_indent = Pt(style.get('left_indent', 0))
        para.paragraph_format.line_spacing = style.get('line_spacing', 0)
        
        # 添加带点号填充的制表位
        # 使用python-docx官方API，避免XML操作冲突
        section = doc.sections[-1]
        
        # 计算正确的制表位位置：考虑左边距和右边距，使页码右对齐且距离右侧边距合适
        # 可用宽度 = 页面宽度 - 左边距 - 右边距
        available_width = section.page_width - section.left_margin - section.right_margin
        # 制表位位置：左边距 + 可用宽度 - 右边距（与右侧边距保持一致）
        tab_position = section.left_margin + available_width - section.right_margin
        
        # 添加右对齐制表位，带点号填充
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        # 先移除所有现有制表位（兼容旧版本python-docx）
        for tab_stop in list(para.paragraph_format.tab_stops):
            para.paragraph_format.tab_stops.remove(tab_stop)
        # 添加新的制表位
        para.paragraph_format.tab_stops.add_tab_stop(
            tab_position,
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.DOTS
        )
        
        # 使用普通文本方式生成目录，确保制表位点号填充正常工作
        
        # 添加标题文本
        text = f"{title['number']} {title['text']}"
        text_run = para.add_run(text)
        text_run.font.name = style.get('font', '')
        text_run._element.rPr.rFonts.set(qn('w:eastAsia'), style.get('font', ''))
        text_run._element.rPr.rFonts.set(qn('w:ascii'), style.get('font', ''))
        text_run._element.rPr.rFonts.set(qn('w:hAnsi'), style.get('font', ''))
        text_run.font.size = Pt(style.get('size', 0))
        # 一级标题强制加粗
        text_run.font.bold = style.get('bold', False) or (title['level'] == 1)
        
        # 添加制表符，触发点号填充
        tab_run = para.add_run('\t')
        
        # 添加页码
        page_config = self.toc_config.get('page_number', {})
        page_run = para.add_run(str(title['page_number']))
        page_run.font.name = page_config.get('font', style.get('font', ''))
        page_run._element.rPr.rFonts.set(qn('w:eastAsia'), page_config.get('font', style.get('font', '')))
        page_run._element.rPr.rFonts.set(qn('w:ascii'), page_config.get('font', style.get('font', '')))
        page_run._element.rPr.rFonts.set(qn('w:hAnsi'), page_config.get('font', style.get('font', '')))
        page_run.font.size = Pt(page_config.get('size', style.get('size', 0)))
        # 一级标题的页码强制加粗
        page_run.font.bold = page_config.get('bold', style.get('bold', False)) or (title['level'] == 1)

class ReferenceManager:
    """动态引用管理器"""
    
    def __init__(self):
        self.references = {}  # {reference_text: number}
        self.reference_numbers = {}  # {number: reference_text}
        self.next_number = 1
    
    def extract_references_from_text(self, text):
        """从文本中提取引用 {{引用内容}} 格式的引用"""
        # 使用字符串方法提取引用
        found_references = []
        ref_start = "{{"
        ref_end = "}}"
        
        pos = 0
        while True:
            # 查找下一个 {{
            start_pos = text.find(ref_start, pos)
            if start_pos == -1:
                break
            
            # 查找对应的 }}
            end_pos = text.find(ref_end, start_pos + len(ref_start))
            if end_pos == -1:
                break
            
            # 提取引用内容
            ref_text = text[start_pos + len(ref_start):end_pos].strip()
            if ref_text:
                found_references.append(ref_text)
                pos = end_pos + len(ref_end)
            else:
                pos = start_pos + 1
        
        return found_references
    
    def process_content_for_references(self, content):
        """处理内容中的引用，替换为动态编号"""
        if not isinstance(content, dict):
            return content
        
        processed_content = {}
        for key, value in content.items():
            if isinstance(value, str):
                # 处理字符串中的引用
                processed_value, found_references = self.replace_reference_markers(value)
                # 记录发现的引用
                for ref_text in found_references:
                    self.add_reference(ref_text)
                processed_content[key] = processed_value
            elif isinstance(value, dict):
                # 递归处理嵌套字典
                processed_content[key] = self.process_content_for_references(value)
            elif isinstance(value, list):
                # 处理列表中的内容
                processed_list = []
                for item in value:
                    if isinstance(item, dict):
                        processed_list.append(self.process_content_for_references(item))
                    elif isinstance(item, str):
                        processed_item, found_references = self.replace_reference_markers(item)
                        for ref_text in found_references:
                            self.add_reference(ref_text)
                        processed_list.append(processed_item)
                    else:
                        processed_list.append(item)
                processed_content[key] = processed_list
            else:
                processed_content[key] = value
        
        return processed_content
    
    def replace_reference_markers(self, text):
        """替换文本中的 {{引用}} 标记为 [数字] 格式，支持连续引用合并为 [n-m] 格式"""
        if not isinstance(text, str):
            return text, []
        
        # 使用字符串方法而不是正则表达式来提取引用
        processed_text = text
        found_references = []
        ref_start = "{{"
        ref_end = "}}"
        
        pos = 0
        while True:
            # 查找下一个 {{
            start_pos = processed_text.find(ref_start, pos)
            if start_pos == -1:
                break
            
            # 查找对应的 }}
            end_pos = processed_text.find(ref_end, start_pos + len(ref_start))
            if end_pos == -1:
                break
            
            # 提取引用内容
            ref_text = processed_text[start_pos + len(ref_start):end_pos].strip()
            if ref_text:
                found_references.append(ref_text)
                # 添加引用到管理器
                number = self.add_reference(ref_text)
                # 替换这个引用标记
                replacement = f"[{number}]"
                processed_text = processed_text[:start_pos] + replacement + processed_text[end_pos + len(ref_end):]
                # 更新pos位置
                pos = start_pos + len(replacement)
            else:
                pos = start_pos + 1
        
        # 合并连续的引用编号为 [n-m] 格式
        processed_text = self._merge_consecutive_references(processed_text)
        
        return processed_text, found_references
    
    def _merge_consecutive_references(self, text):
        """将连续的引用编号 [n][m][k] 合并为 [n-m-k] 或 [n-k] 格式"""
        import re
        
        # 查找所有引用标记 [数字]
        pattern = r'(\[\d+\])'
        matches = list(re.finditer(pattern, text))
        
        if len(matches) < 2:
            return text
        
        # 提取所有引用编号和位置
        ref_data = []
        for match in matches:
            ref_data.append({
                'number': int(match.group(1)[1:-1]),
                'start': match.start(),
                'end': match.end(),
                'text': match.group(0)
            })
        
        # 找出连续编号的位置组
        ref_data = sorted(ref_data, key=lambda x: x['start'])
        
        # 构建合并任务列表
        merge_tasks = []
        i = 0
        while i < len(ref_data):
            j = i
            group = [ref_data[i]]
            while j + 1 < len(ref_data):
                # 检查是否连续（编号连续且位置连续）
                curr = ref_data[j]
                next_item = ref_data[j + 1]
                if (next_item['number'] == curr['number'] + 1 and 
                    next_item['start'] == curr['end']):
                    group.append(next_item)
                    j += 1
                else:
                    break
            
            if len(group) >= 2:
                merge_tasks.append(group)
            i = j + 1
        
        if not merge_tasks:
            return text
        
        # 从后向前执行替换
        result = text
        for group in reversed(merge_tasks):
            # 计算合并后的替换文本
            numbers = [item['number'] for item in group]
            if len(numbers) == 2:
                merged_ref = f"[{numbers[0]}-{numbers[1]}]"
            else:
                merged_ref = f"[{numbers[0]}-{numbers[-1]}]"
            
            # 从后向前替换，确保位置不受影响
            for item in reversed(group):
                result = result[:item['start']] + '___PLACEHOLDER___' + result[item['end']:]
            
            # 替换占位符为合并后的引用
            result = result.replace('___PLACEHOLDER___', merged_ref, 1)
            # 清除其他占位符
            result = result.replace('___PLACEHOLDER___', '')
        
        return result
    
    def _find_consecutive_groups(self, numbers):
        """找出数字列表中的连续序列"""
        if not numbers:
            return []
        
        # 按数字大小排序
        sorted_numbers = sorted(numbers)
        groups = []
        current_group = [sorted_numbers[0]]
        
        for i in range(1, len(sorted_numbers)):
            if sorted_numbers[i] == sorted_numbers[i-1] + 1:
                # 连续，加入当前组
                current_group.append(sorted_numbers[i])
            else:
                # 不连续，保存当前组并开始新组
                if len(current_group) >= 2:
                    groups.append(current_group)
                current_group = [sorted_numbers[i]]
        
        # 处理最后一个组
        if len(current_group) >= 2:
            groups.append(current_group)
        
        return groups
    
    def add_reference(self, ref_text):
        """添加引用，返回分配的序号"""
        ref_text = ref_text.strip()
        if not ref_text:
            return None
        
        # 如果引用已存在，返回现有序号
        if ref_text in self.references:
            return self.references[ref_text]
        
        # 添加新引用
        number = self.next_number
        self.references[ref_text] = number
        self.reference_numbers[number] = ref_text
        self.next_number += 1
        
        return number
    
    def get_all_references(self):
        """获取所有引用，按序号排序"""
        return [(number, text) for number, text in sorted(self.reference_numbers.items())]
    
    def clear(self):
        """清空所有引用"""
        self.references.clear()
        self.reference_numbers.clear()
        self.next_number = 1

# 配置参数 - 从YAML文件读取
CONFIG = {
    # 文件配置
    'default_config_file': os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Demo_Config.yaml'),
    'default_output_filename': 'Format_Reference.docx'
}

# ==================== 功能函数区域 ====================

def set_style_properties(style, font_name, font_size, bold=False, align=None, 
                        first_line_indent=None, space_before=None, space_after=None, 
                        line_spacing=1.0, hanging_indent=None, left_indent=None, color=None):
    """设置样式属性"""
    font = style.font
    font.name = font_name
    font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    font._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    font.size = Pt(font_size)
    font.bold = bold
    
    # 设置字体颜色，默认为黑色
    if color is None:
        # 黑色RGB值为RGBColor(0, 0, 0)
        font.color.rgb = RGBColor(0, 0, 0)
    else:
        font.color.rgb = color
    
    pf = style.paragraph_format
    if align:
        pf.alignment = align
    if first_line_indent:
        pf.first_line_indent = Pt(first_line_indent)
    if space_before:
        pf.space_before = Pt(space_before)
    if space_after:
        pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if hanging_indent:
        pf.hanging_indent = Pt(hanging_indent)
    if left_indent:
        pf.left_indent = Pt(left_indent)


def set_run_font(run, font_name, font_size, bold=False, color=None):
    """设置Run对象的字体属性"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    
    # 设置字体颜色，默认为黑色
    if color is None:
        # 黑色RGB值为RGBColor(0, 0, 0)
        run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run.font.color.rgb = color


def create_xml_element(tag_name, attributes=None):
    """创建XML元素"""
    if attributes is None:
        attributes = {}
    
    # 构建XML字符串
    attr_str = ""
    for key, value in attributes.items():
        attr_str += f" {key}=\"{value}\""
    
    xml_str = f'<w:{tag_name} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"{attr_str}/>'
    return parse_xml(xml_str)


def add_run_to_hyperlink(hyperlink, text, font_name, font_size, bold=False):
    """向超链接添加运行元素"""
    # 创建运行元素
    run = create_xml_element('r')
    hyperlink.append(run)
    
    # 设置运行属性
    rPr = create_xml_element('rPr')
    run.append(rPr)
    
    # 设置字体
    rFonts = create_xml_element('rFonts', {
        'w:eastAsia': font_name,
        'w:ascii': font_name,
        'w:hAnsi': font_name
    })
    rPr.append(rFonts)
    
    # 设置字体大小
    font_size_half = int(font_size * 2)
    sz = create_xml_element('sz', {'w:val': str(font_size_half)})
    rPr.append(sz)
    
    # 设置粗体
    if bold:
        bold_elem = create_xml_element('b')
        rPr.append(bold_elem)
    
    # 添加文本
    text_elem = create_xml_element('t', {'xml:space': 'preserve'})
    text_elem.text = text
    run.append(text_elem)
    
    return run

def load_config_file(config_file):
    """加载YAML配置文件"""
    if not os.path.exists(config_file):
        raise FileNotFoundError(f"配置文件未找到")
    
    with open(config_file, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)

def load_multi_config_file(config_file):
    """加载多文件配置文件，支持基础配置和章节配置的合并"""
    if not os.path.exists(config_file):
        raise FileNotFoundError(f"配置文件未找到")
    
    # 获取配置文件的目录路径，用于解析相对路径
    config_dir = os.path.dirname(os.path.abspath(config_file))
    
    with open(config_file, 'r', encoding='utf-8') as f:
        master_config = yaml.safe_load(f)
    
    merged_config = {}
    
    # 创建引用管理器实例
    ref_manager = ReferenceManager()
    
    # 加载基础配置文件
    if "import_base_config" in master_config:
        base_config_file = master_config["import_base_config"]
        # 如果是相对路径，转换为相对于主配置文件的绝对路径
        if not os.path.isabs(base_config_file):
            base_config_file = os.path.join(config_dir, base_config_file)
            
        if os.path.exists(base_config_file):
            with open(base_config_file, 'r', encoding='utf-8') as f:
                base_config = yaml.safe_load(f)
                # 处理基础配置中的引用
                base_config = ref_manager.process_content_for_references(base_config)
                merged_config.update(base_config)
    
    # 加载组件配置文件（封面页、目录页等）
    if "component_configs" in master_config:
        for component_file in master_config["component_configs"]:
            # 如果是相对路径，转换为相对于主配置文件的绝对路径
            if not os.path.isabs(component_file):
                component_file = os.path.join(config_dir, component_file)
                
            if os.path.exists(component_file):
                with open(component_file, 'r', encoding='utf-8') as f:
                    component_config = yaml.safe_load(f)
                    # 处理组件配置中的引用
                    component_config = ref_manager.process_content_for_references(component_config)
                    # 将组件配置合并到主配置中
                    merged_config.update(component_config)
    
    # 保存配置文件目录，用于后续解析相对路径
    merged_config['config_dir'] = config_dir
    
    # 加载章节配置文件
    if "chapter_configs" in master_config:
        sections = []
        for chapter_file in master_config["chapter_configs"]:
            # 如果是相对路径，转换为相对于主配置文件的绝对路径
            if not os.path.isabs(chapter_file):
                chapter_file = os.path.join(config_dir, chapter_file)
                
            if os.path.exists(chapter_file):
                with open(chapter_file, 'r', encoding='utf-8') as f:
                    chapter_config = yaml.safe_load(f)
                    # 处理章节配置中的引用
                    chapter_config = ref_manager.process_content_for_references(chapter_config)
                    if "sections" in chapter_config:
                        sections.extend(chapter_config["sections"])
                    elif "title" in chapter_config and "author" in chapter_config:
                        # 单独处理文档基本信息
                        if "document_info" not in merged_config:
                            merged_config["document_info"] = {}
                        merged_config["document_info"].update({
                            "title": chapter_config["title"]["text"] if "text" in chapter_config["title"] else merged_config.get("document_info", {}).get("title"),
                            "author": chapter_config["author"]["text"] if "text" in chapter_config["author"] else merged_config.get("document_info", {}).get("author")
                        })
        
        # 添加合并后的章节内容
        if sections:
            if "content_structure" not in merged_config:
                merged_config["content_structure"] = {}
            merged_config["content_structure"]["sections"] = sections
    
    # 合并主配置文件中的其他信息
    for key, value in master_config.items():
        if key not in ["import_base_config", "chapter_configs", "content_structure"]:
            # 处理主配置文件中的引用
            merged_config[key] = ref_manager.process_content_for_references(value)
    
    # 将引用管理器实例添加到配置中，以便后续使用
    merged_config["_reference_manager"] = ref_manager
    
    return merged_config

def get_alignment_constant(alignment_str):
    """将字符串对齐方式转换为常量"""
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    return alignment_map.get(alignment_str, WD_ALIGN_PARAGRAPH.LEFT)

def add_page_numbers(doc, abstract_end_paragraph=None, config=None):
    """为文档添加页码，封面不显示页码，目录和摘要使用大写罗马数字，正文使用阿拉伯数字"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    
    # 从配置文件获取页码样式
    page_num_config = config.get('page_number', {}) if config else {}
    
    # 获取节(section)对象
    sections = doc.sections
    
    # 完全重置所有页脚
    for section in sections:
        footer = section.footer
        # 清除所有现有页脚内容
        for para in footer.paragraphs:
            para.clear()
        # 禁用与前一节相同的页脚
        footer.is_linked_to_previous = False
    
    # 为所有页设置阿拉伯数字页码
    for idx, section in enumerate(sections):
        footer = section.footer
        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        # 使用Word域代码生成阿拉伯数字页码
        fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        instrText = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
        fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.name = page_num_config.get('font', 'Times New Roman')
        run.font.size = Pt(page_num_config.get('size', 12))
        # 设置页码起始编号
        sectPr = section._sectPr
        if sectPr is not None:
            if idx == 0:  # 封面页
                pgNumType = parse_xml(r'<w:pgNumType w:fmt="decimal" w:start="1" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            else:  # 其他页继续编号
                pgNumType = parse_xml(r'<w:pgNumType w:fmt="decimal" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            sectPr.append(pgNumType)
    else:
        # 处理文档只有3个或更少section的情况
        for i in range(len(sections)):
            if i >= 1:  # 跳过封面页
                footer = sections[i].footer
                footer_para = footer.add_paragraph()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                footer_run = footer_para.add_run()
                # 使用Word域代码生成阿拉伯数字页码
                fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                instrText = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE \* MERGEFORMAT </w:instrText>')
                fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                footer_run._r.append(fldChar1)
                footer_run._r.append(instrText)
                footer_run._r.append(fldChar2)
                footer_run.font.name = page_num_config.get('font', 'Times New Roman')
                footer_run.font.size = Pt(page_num_config.get('size', 12))
def add_paragraph_with_style(doc, text, style_name, alignment=None, title_collector=None, bookmark_id=None):
    """添加带样式的段落，支持软换行和参考文献上标格式，可为标题添加书签"""
    # 为标题段落添加书签
    is_title = False
    if title_collector and (style_name in ['\u200D一级标题', '\u200e二级标题', '\u200F三级标题'] or 
                            '一级标题' in style_name or '二级标题' in style_name or '三级标题' in style_name):
        is_title = True
    
    # 如果是标题，先添加书签占位符
    if is_title and bookmark_id is None:
        bookmark_id = len(doc.paragraphs)
    
    # 添加段落
    para = doc.add_paragraph()
    
    # 如果是标题，添加书签
    if is_title and bookmark_id is not None:
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn
        
        # 获取段落属性
        pPr = para._element.get_or_add_pPr()
        
        # 创建书签开始标记
        bookmark_start = parse_xml(f'<w:bookmarkStart w:id="{bookmark_id}" w:name="toc_{bookmark_id}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        pPr.insert(0, bookmark_start)
        
        # 处理软换行：将 \n 转换为 Word 中的换行
        if '\n' in text:
            lines = text.split('\n')
            for i, line in enumerate(lines):
                if line.strip():  # 跳过空行
                    if i > 0:  # 不是第一行
                        para.add_run('\n')
                    add_text_with_sub_sup(para, line.strip())
        else:
            add_text_with_sub_sup(para, text)
        
        # 创建书签结束标记
        bookmark_end = parse_xml(f'<w:bookmarkEnd w:id="{bookmark_id}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        pPr.append(bookmark_end)
    else:
        # 非标题段落，正常处理
        # 处理软换行：将 \n 转换为 Word 中的换行
        if '\n' in text:
            lines = text.split('\n')
            for i, line in enumerate(lines):
                if line.strip():  # 跳过空行
                    if i > 0:  # 不是第一行
                        para.add_run('\n')
                    add_text_with_sub_sup(para, line.strip())
        else:
            add_text_with_sub_sup(para, text)
    
    if alignment:
        para.alignment = alignment
    
    # 设置段落样式
    para.style = style_name
    
    # 如果提供了标题收集器且是标题样式，则收集标题信息
    if title_collector and is_title:
        # 计算当前段落的页码（简化处理，实际应该根据分页符计算）
        current_page = 1  # 初始假设
        paragraph_index = len(doc.paragraphs) - 1
        title_collector.add_title(text, style_name, paragraph_index, current_page)
    
    return para

def add_text_with_superscript_refs(para, text):
    """添加文本，将参考文献标记 [数字] 格式设置为上标"""
    # 使用正则表达式匹配 [数字] 格式的参考文献标记
    pattern = r'\[(\d+)\]'
    
    # 查找所有匹配项
    last_end = 0
    for match in re.finditer(pattern, text):
        # 添加匹配前的普通文本
        if match.start() > last_end:
            para.add_run(text[last_end:match.start()])
        
        # 添加上标的参考文献标记
        ref_run = para.add_run(f"[{match.group(1)}]")
        ref_run.font.superscript = True
        
        last_end = match.end()
    
    # 添加剩余的文本
    if last_end < len(text):
        para.add_run(text[last_end:])


def add_text_with_sub_sup(para, text):
    """添加文本，处理 <sub> 和 <sup> 标签实现上下标，同时支持参考文献上标标记和斜体

    语法说明：
    - <sub>内容</sub>  将内容设置为下标，如 CO<sub>2</sub> 显示为 CO₂
    - <sup>内容</sup>  将内容设置为上标，如 H<sup>+</sup> 显示为 H⁺
    - <i>内容</i>  将内容设置为斜体，如 <i>italic</i> 显示为 italic
    - [数字] 自动设置为上标，如 [1] 显示为参考文献上标格式

    示例：
    - 水分子：H<sub>2</sub>O
    - 氢离子：H<sup>+</sup>
    - 硫酸根：SO<sub>4</sub><sup>2-</sup>
    - 引用标注：H<sub>2</sub>O 是生命之源[1]
    - 斜体文本：<i>这是斜体文本</i>
    """
    import re

    sub_pattern = re.compile(r'<sub>(.*?)</sub>')
    sup_pattern = re.compile(r'<sup>(.*?)</sup>')
    italic_pattern = re.compile(r'<i>(.*?)</i>')
    ref_pattern = re.compile(r'\[(\d+)\]')

    last_end = 0
    matches = []

    sub_matches = [(m, 'sub') for m in sub_pattern.finditer(text)]
    sup_matches = [(m, 'sup') for m in sup_pattern.finditer(text)]
    italic_matches = [(m, 'italic') for m in italic_pattern.finditer(text)]
    ref_matches = [(m, 'ref') for m in ref_pattern.finditer(text)]

    all_matches = sub_matches + sup_matches + italic_matches + ref_matches
    all_matches.sort(key=lambda x: x[0].start())

    for match, match_type in all_matches:
        start, end = match.span()
        if start > last_end:
            para.add_run(text[last_end:start])

        if match_type == 'sub':
            sub_run = para.add_run(match.group(1))
            sub_run.font.subscript = True
        elif match_type == 'sup':
            sup_run = para.add_run(match.group(1))
            sup_run.font.superscript = True
        elif match_type == 'italic':
            italic_run = para.add_run(match.group(1))
            italic_run.font.italic = True
        elif match_type == 'ref':
            ref_run = para.add_run(f"[{match.group(1)}]")
            ref_run.font.superscript = True

        last_end = end

    if last_end < len(text):
        para.add_run(text[last_end:])


def count_subscript_superscript_in_doc(doc):
    """统计文档中所有段落的上下标数量
    
    返回: (下标数量, 上标数量)
    """
    global subscript_count, superscript_count
    subscript_count = 0
    superscript_count = 0
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.subscript:
                subscript_count += 1
            if run.font.superscript:
                superscript_count += 1
    
    return subscript_count, superscript_count


def add_reference_paragraph(doc, formatted_ref, has_chinese, config=None):
    """添加参考文献段落，根据语言设置不同字体"""
    para = doc.add_paragraph()
    para.style = '\u206F参考文献'
    run = para.add_run(formatted_ref)
    
    # 从配置文件获取参考文献样式
    if config:
        ref_config = config.get('text_styles', {}).get('reference', {})
        font = ref_config.get('font', '宋体')
        english_font = ref_config.get('english_font', 'Times New Roman')
        size = ref_config.get('size', 10.5)
        
        if has_chinese:
            # 中文引用：使用配置的字体
            set_run_font(run, font, size)
        else:
            # 英文引用：使用配置的西文字体
            set_run_font(run, english_font, size)
    else:
        # 默认样式
        if has_chinese:
            # 中文引用：使用宋体五号
            set_run_font(run, '宋体', 10.5)
        else:
            # 英文引用：使用Times New Roman
            set_run_font(run, 'Times New Roman', 10.5)
    
    return para

def add_numbered_list(doc, items, level=1, config=None):
    """添加带编号的列表"""
    # 从配置文件获取列表设置
    list_config = config.get('list', {})
    font_config = config.get('font', {})
    
    for i, item in enumerate(items, 1):
        para = doc.add_paragraph()
        if level == 1:
            number = list_config.get('main_prefix', "").format(i=i)
        else:
            circles = list_config.get('circles', ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨", "⑩",
                       "⑪", "⑫", "⑬", "⑭", "⑮", "⑯", "⑰", "⑱", "⑲", "⑳"])
            number = circles[i-1] if i <= len(circles) else f'({i})'
        
        run = para.add_run(f"{number} {item}")
        run.font.name = font_config.get('default', '')
        run.font.size = Pt(font_config.get('default_size', 0))
        
        if level == 1:
            para.paragraph_format.space_before = Pt(list_config.get('main_spacing', 6))
        else:
            para.paragraph_format.left_indent = Pt(list_config.get('sub_left_indent', 24))
            para.paragraph_format.first_line_indent = Pt(list_config.get('sub_first_indent', -24))

def get_image_info(image_path):
    """获取图片信息"""
    from PIL import Image
    try:
        with Image.open(image_path) as img:
            return img.size, img.size[1] / img.size[0]
    except Exception as e:
        return (200, 150), 0.75

def calculate_image_dimensions(image_paths, images_per_row=1, config=None):
    """统一计算图片尺寸，确保高度对齐且填满宽度"""
    image_config = config.get('image', {})
    spacing = image_config.get('spacing', 5)
    
    # 计算可用宽度和高度
    page_width = image_config.get('page_width', 595)
    page_height = image_config.get('page_height', 842)
    page_margins = image_config.get('page_margins', (110, 110))
    available_width = page_width - sum(page_margins)
    margins = page_margins[0]
    title_space = 200
    available_height = max(page_height - margins * 2 - title_space, 300)
    
    image_infos = [get_image_info(path) for path in image_paths]
    
    if images_per_row == 1:
        width = available_width
        width = max(image_config.get('min_width', 50), min(width, image_config.get('max_width', 5000)))
        _, aspect_ratio = image_infos[0]
        height = width * aspect_ratio
        
        if height > available_height:
            scale = available_height / height
            width *= scale
            height = available_height
            
        return [(width, height)], height, image_infos
    
    else:
        total_spacing = spacing * (images_per_row - 1)
        width_per_image = (available_width - total_spacing) / images_per_row
        width_per_image = max(image_config.get('min_width', 50), 
                             min(width_per_image, image_config.get('max_width', 5000)))
        
        dimensions = []
        max_height = 0
        for size, aspect_ratio in image_infos:
            height = width_per_image * aspect_ratio
            max_height = max(max_height, height)
            dimensions.append((width_per_image, height))
        
        if max_height > available_height:
            scale = available_height / max_height
            dimensions = [(w * scale, h * scale) for w, h in dimensions]
            max_height = available_height
        
        return dimensions, max_height, image_infos



def add_adaptive_images(doc, image_paths, captions=None, config=None):
    """自适应图片排版"""
    global image_counter, small_image_count, large_image_count
    
    if not image_paths:
        return
    
    num_images = len(image_paths)
    large_image_count += 1
    small_image_count += num_images
    image_config = config.get('image', {})
    spacing = image_config.get('spacing', 5)
    
    # 处理图片路径
    config_dir = config.get('config_dir', '')
    resolved_image_paths = []
    for image_path in image_paths:
        if image_path and not os.path.isabs(image_path):
            resolved_path = os.path.join(config_dir, image_path)
            resolved_image_paths.append(resolved_path)
        else:
            resolved_image_paths.append(image_path)
    
    try:
        current_image_number = image_counter
        
        if num_images == 1:
            target_width, target_height = calculate_image_dimensions(resolved_image_paths, 1, config)[0][0]
            
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(resolved_image_paths[0], width=Pt(target_width))
            
            # 设置图像与图题之间的间距
            fig_config = image_config.get('figure', {})
            para.paragraph_format.space_after = Pt(image_config.get('caption_spacing', 3))
            
            if captions:
                if captions[0] and captions[0].strip():
                    caption_text = f"图{current_image_number} {captions[0].strip()}"
                else:
                    caption_text = f"图{current_image_number}"
                caption_para = doc.add_paragraph(caption_text, style='\u200C图题')
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            image_counter += 1
        
        else:  # 多张图片情况：2-4张
            if num_images == 2:
                target_dimensions, max_height, _ = calculate_image_dimensions(resolved_image_paths, 2, config)
                
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_after = Pt(image_config.get('caption_spacing', 3))
                
                for i, (width, height) in enumerate(target_dimensions):
                    run = para.add_run()
                    run.add_picture(resolved_image_paths[i], width=Pt(width), height=Pt(max_height))
                    if i == 0:
                        spacing_run = para.add_run(" " * 3)
                        spacing_run.font.size = Pt(spacing // 2)
            
            elif num_images == 3:
                first_row_dimensions, first_row_max_height, _ = calculate_image_dimensions(resolved_image_paths[:2], 2, config)
                
                para1 = doc.add_paragraph()
                para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para1.paragraph_format.space_before = Pt(0)
                
                for i in range(2):
                    width, _ = first_row_dimensions[i]
                    run = para1.add_run()
                    run.add_picture(resolved_image_paths[i], width=Pt(width), height=Pt(first_row_max_height))
                    if i == 0:
                        spacing_run = para1.add_run(" " * 3)
                        spacing_run.font.size = Pt(spacing // 2)
                
                para1.paragraph_format.space_after = Pt(image_config.get('between_images_spacing', 1))
                
                first_row_width = first_row_dimensions[0][0] * 2 + spacing
                third_width = max(image_config.get('min_width', 50), 
                                min(first_row_width, image_config.get('max_width', 5000)))
                _, third_aspect_ratio = get_image_info(resolved_image_paths[2])
                third_height = third_width * third_aspect_ratio
                
                para2 = doc.add_paragraph()
                para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para2.paragraph_format.space_before = Pt(0)
                para2.paragraph_format.space_after = Pt(image_config.get('between_images_spacing', 1))
                para2.add_run().add_picture(resolved_image_paths[2], width=Pt(third_width), height=Pt(third_height))
            
            elif num_images == 4:
                first_row_dimensions, first_row_max_height, _ = calculate_image_dimensions(resolved_image_paths[:2], 2, config)
                
                para1 = doc.add_paragraph()
                para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for i in range(2):
                    width, _ = first_row_dimensions[i]
                    run = para1.add_run()
                    run.add_picture(resolved_image_paths[i], width=Pt(width), height=Pt(first_row_max_height))
                    if i == 0:
                        spacing_run = para1.add_run(" " * 3)
                        spacing_run.font.size = Pt(spacing // 2)
                
                para1.paragraph_format.space_after = Pt(image_config.get('between_images_spacing', 1))
                
                first_row_width = first_row_dimensions[0][0] * 2 + spacing
                target_width2 = max(image_config.get('min_width', 50), 
                                  min((first_row_width - spacing) / 2, image_config.get('max_width', 5000)))
                
                second_row_max_height = 0
                for i in range(2):
                    _, aspect_ratio = get_image_info(resolved_image_paths[2 + i])
                    height = target_width2 * aspect_ratio
                    second_row_max_height = max(second_row_max_height, height)
                
                para2 = doc.add_paragraph()
                para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para2.paragraph_format.space_before = Pt(0)
                para2.paragraph_format.space_after = Pt(image_config.get('between_images_spacing', 1))
                
                for i in range(2):
                    run = para2.add_run()
                    run.add_picture(resolved_image_paths[2 + i], width=Pt(target_width2), height=Pt(second_row_max_height))
                    if i == 0:
                        spacing_run = para2.add_run(" " * 3)
                        spacing_run.font.size = Pt(spacing // 2)
            
            # 生成统一的图题
            if captions:
                caption_para = doc.add_paragraph(style='\u200C图题')
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fig_config = image_config.get('figure', {})
                caption_para.paragraph_format.space_before = Pt(fig_config.get('space_before', 0))
                caption_para.paragraph_format.space_after = Pt(fig_config.get('space_after', 6))
                
                caption_parts = [f"图{current_image_number}"]
                for i in range(num_images):
                    sub_letter = chr(ord('a') + i)
                    if i < len(captions):
                        original_caption = captions[i]
                        sub_caption = f"({sub_letter}) {original_caption}"
                    else:
                        sub_caption = f"({sub_letter}) 图片"
                    caption_parts.append(sub_caption)
                
                full_caption = " ".join(caption_parts)
                caption_para.add_run(full_caption)
            
            image_counter += 1
        
    except Exception as e:
        import traceback
        traceback.print_exc()

# ==================== 表格功能区域 ====================

def set_cell_border(cell, **kwargs):
    """
    设置表格单元格边框
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # 设置边框
    tcBorders = OxmlElement('w:tcBorders')
    tcPr.append(tcBorders)
    
    # 遍历所有边框类型
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            
            # 设置边框样式和宽度
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_table_borders(table):
    """移除表格所有默认边框"""
    for row in table.rows:
        for cell in row.cells:
            # 清除所有边框
            set_cell_border(cell, 
                          top={"sz": 0, "val": "nil"},
                          bottom={"sz": 0, "val": "nil"},
                          start={"sz": 0, "val": "nil"},
                          end={"sz": 0, "val": "nil"},
                          insideH={"sz": 0, "val": "nil"},
                          insideV={"sz": 0, "val": "nil"})

def add_table_caption(doc, caption_text, table_number, config=None):
    """添加表格标题 - 表序和表题居中排版，表序居左与表题间空一汉字间距"""
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 从配置文件获取表格样式
    table_config = config.get('table', {}) if config else {}
    table_num_config = table_config.get('table_num', {})
    caption_config = table_config.get('caption', {})
    
    # 表序：居左，与表题间空一个汉字间距
    table_num_run = caption_para.add_run(f"表{table_number}")
    table_num_run.font.name = table_num_config.get('font', '宋体')
    table_num_run.font.size = Pt(table_num_config.get('size', 10.5))
    table_num_run.font.bold = table_num_config.get('bold', True)
    if 'Times' not in table_num_config.get('font', ''):
        table_num_run._element.rPr.rFonts.set(qn('w:eastAsia'), table_num_config.get('font', '宋体'))
    
    # 空一汉字字符间距
    space_run = caption_para.add_run("　")  # 全角空格
    
    # 表题：不加任何标点符号
    caption_run = caption_para.add_run(caption_text)
    caption_run.font.name = caption_config.get('font', '宋体')
    caption_run.font.size = Pt(caption_config.get('size', 10.5))
    caption_run.font.bold = caption_config.get('bold', True)
    if 'Times' not in caption_config.get('font', ''):
        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), caption_config.get('font', '宋体'))
    
    # 设置表题与表格之间的间距
    caption_spacing = table_config.get('caption_spacing', 0)
    caption_para.paragraph_format.space_after = Pt(caption_spacing)
    
    return caption_para

def apply_mixed_font_to_cell(cell, text, is_header=False, config=None):
    """为单元格应用混合字体：中文使用宋体，英文使用Times New Roman
    
    Args:
        cell: 单元格对象
        text: 文本内容
        is_header: 是否为表头，用于设置加粗格式
        config: 配置参数，用于获取字体设置
    """
    # 清除原有内容
    cell.text = ""
    
    # 创建新的段落
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 从配置中获取字体设置
    table_config = config.get('table', {}) if config else {}
    header_config = table_config.get('header', {})
    
    # 设置默认字体和大小
    chinese_font = header_config.get('font', '宋体')
    english_font = 'Times New Roman'
    font_size = header_config.get('size', 10.5)
    
    # 遍历文本中的每个字符，分别设置字体
    i = 0
    while i < len(text):
        char = text[i]
        
        # 检测是否为英文字符
        if char.isalpha() and ord(char) < 128:  # ASCII字母
            # 找到连续英文字符的结束位置
            j = i
            while j < len(text) and text[j].isalpha() and ord(text[j]) < 128:
                j += 1
            english_text = text[i:j]
            
            # 添加英文字符，Times New Roman字体
            run = para.add_run(english_text)
            run.font.name = english_font
            run.font.size = Pt(font_size)
            
            # 设置英文字体的东亚字体映射
            run._element.rPr.rFonts.set(qn('w:eastAsia'), english_font)
            
            i = j
        else:
            # 中文字符或符号，使用中文字体
            run = para.add_run(char)
            run.font.name = chinese_font
            run.font.size = Pt(font_size)
            
            # 设置中文字体的东亚字体映射
            run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
            
            i += 1

def create_three_line_table(doc, data, caption_text=None, config=None):
    """
    创建三线格表格
    
    Args:
        doc: Word文档对象
        data: 表格数据，格式为[[headers], [row1_data], [row2_data], ...]
        caption_text: 表格标题文字
        config: 配置参数
    
    Returns:
        表格对象
    """
    global table_counter
    
    if not data or len(data) < 2:
    
        return None
    
    rows = len(data)
    cols = len(data[0])
    
    # 使用自动表格编号
    current_table_number = table_counter
    table_counter += 1
    
    # 添加表格标题（位于表格上方）
    table_config = config.get('table', {}) if config else {}
    if caption_text and table_config.get('caption_above', True):
        add_table_caption(doc, caption_text, current_table_number, config)
    
    # 创建表格
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表格样式
    table.style = 'Table Grid'
    
    # 移除所有默认边框
    set_table_borders(table)
    
    # 填充数据，应用混合字体
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            is_header = (row_idx == 0)
            apply_mixed_font_to_cell(cell, str(cell_data), is_header=is_header, config=config)
    
    # 应用三线格格式
    apply_three_line_formatting(table)
    
    return table

def apply_three_line_formatting(table):
    """
    应用三线格格式
    - 顶部线：1pt
    - 底部线：0.5pt  
    - 表头底线：0.5pt
    """
    rows = table.rows
    total_rows = len(rows)
    
    if total_rows < 2:
        return
    
    # 第一行（表头）- 设置顶部和底部边框
    header_row = rows[0]
    for cell in header_row.cells:
        set_cell_border(cell,
                       top={"sz": 16, "val": "single", "color": "000000"},  # 1pt = 16/16
                       bottom={"sz": 8, "val": "single", "color": "000000"}) # 0.5pt = 8/16
    
    # 最后一行（如果有数据行）- 设置底部边框
    if total_rows > 1:
        last_row = rows[-1]
        for cell in last_row.cells:
            set_cell_border(cell,
                           bottom={"sz": 8, "val": "single", "color": "000000"}) # 0.5pt



class CoverPageGenerator:
    """封面页生成器"""
    
    def __init__(self, cover_config, config_dir=None):
        self.cover_config = cover_config
        self.config_dir = config_dir
        # 从配置文件加载样式
        self.cover_styles = {
            'school_name': cover_config.get('school_name', {}),
            'title_chinese': cover_config.get('title', {}),
            'title_english': cover_config.get('title', {}),
            'info_label': cover_config.get('info_label', cover_config.get('label_style', {})),
            'info_content': cover_config.get('info_content', cover_config.get('content_style', {})),
            'footer': cover_config.get('footer', {})
        }
    
    def generate_chinese_cover(self, doc, doc_info):
        """生成中文封面页 - 按照标准上海交通大学学位论文格式"""
        
        # 1. 第一行是图标
        cover_image_config = self.cover_config.get('cover_image', {})
        cover_image_path = cover_image_config.get('path', "")
        
        # 处理相对路径
        if cover_image_path and self.config_dir and not os.path.isabs(cover_image_path):
            cover_image_path = os.path.join(self.config_dir, cover_image_path)
            cover_image_path = os.path.normpath(cover_image_path)
        
        if os.path.exists(cover_image_path):
            img_para = doc.add_paragraph()
            img_para.alignment = get_alignment_constant(cover_image_config.get('alignment', ''))
            run = img_para.add_run()
            picture = run.add_picture(
                cover_image_path, 
                width=Cm(cover_image_config.get('width_cm', 3.13)), 
                height=Cm(cover_image_config.get('height_cm', 3.13))
            )
        
        # 2. 以五号字体空一行
        empty_line_config = self.cover_config.get('empty_line', {})
        empty_para1 = doc.add_paragraph()
        empty_para1.paragraph_format.space_after = Pt(empty_line_config.get('space_after', 10.5))
        
        # 3. 学校名称
        school_name_config = self.cover_config.get('school_name', {})
        school_para = doc.add_paragraph()
        school_run = school_para.add_run(school_name_config.get('text', ''))
        school_run.font.name = school_name_config.get('font', '')
        school_run._element.rPr.rFonts.set(qn('w:eastAsia'), school_name_config.get('font', ''))
        school_run._element.rPr.rFonts.set(qn('w:ascii'), school_name_config.get('font', ''))
        school_run._element.rPr.rFonts.set(qn('w:hAnsi'), school_name_config.get('font', ''))
        school_run.font.size = Pt(school_name_config.get('size', 0))
        school_run.font.bold = school_name_config.get('bold', True)
        school_para.alignment = get_alignment_constant(school_name_config.get('alignment', ''))
        school_para.paragraph_format.space_after = Pt(school_name_config.get('space_after', 36))
        
        # 4. 论文标题
        title_config = self.cover_config.get('title', {})
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(doc_info.get('title', ''))
        title_run.font.name = title_config.get('font', '')
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_config.get('font', ''))
        title_run._element.rPr.rFonts.set(qn('w:ascii'), title_config.get('font', ''))
        title_run._element.rPr.rFonts.set(qn('w:hAnsi'), title_config.get('font', ''))
        title_run.font.size = Pt(title_config.get('size', 0))
        title_run.font.bold = title_config.get('bold', False)
        title_para.alignment = get_alignment_constant(title_config.get('alignment', ''))
        title_para.paragraph_format.line_spacing = title_config.get('line_spacing', 0)
        title_para.paragraph_format.space_after = Pt(title_config.get('space_after', 0))
        
        # 空四行
        for _ in range(4):
            doc.add_paragraph()
        
        # 5. 学生信息部分
        info_label_config = self.cover_config.get('info_label', {})
        info_content_config = self.cover_config.get('info_content', {})
        
        # 从配置文件读取info_items
        info_items_config = self.cover_config.get('info_items', [])
        
        info_items = []
        for item in info_items_config:
            label = item.get('label', '')
            field = item.get('field', '')
            if field in doc_info and doc_info[field]:
                info_items.append((label, doc_info[field]))
        
        # 如果没有提供这些信息，使用空列表
        if not info_items:
            info_items = []
        
        for label, value in info_items:
            info_para = doc.add_paragraph()
            info_para.alignment = get_alignment_constant(info_label_config.get('alignment', ''))
            
            # 设置左缩进
            info_para.paragraph_format.left_indent = Pt(info_label_config.get('left_indent', 0))
            
            # 冒号前：配置字体
            label_run = info_para.add_run(label)
            label_run.font.name = info_label_config.get('font', '')
            label_run._element.rPr.rFonts.set(qn('w:eastAsia'), info_label_config.get('font', ''))
            label_run._element.rPr.rFonts.set(qn('w:ascii'), info_label_config.get('font', ''))
            label_run._element.rPr.rFonts.set(qn('w:hAnsi'), info_label_config.get('font', ''))
            label_run.font.size = Pt(info_label_config.get('size', 0))
            label_run.font.bold = info_label_config.get('bold', False)
            
            # 冒号后：配置字体
            value_run = info_para.add_run(value)
            value_run.font.name = info_content_config.get('font', '')
            value_run._element.rPr.rFonts.set(qn('w:eastAsia'), info_content_config.get('font', ''))
            value_run._element.rPr.rFonts.set(qn('w:ascii'), info_content_config.get('font', ''))
            value_run._element.rPr.rFonts.set(qn('w:hAnsi'), info_content_config.get('font', ''))
            value_run.font.size = Pt(info_content_config.get('size', 0))
            value_run.font.bold = info_content_config.get('bold', False)
            
            info_para.paragraph_format.space_after = Pt(info_label_config.get('space_after', 18))
        
        # 6. 换3行，写下日期
        for _ in range(3):
            doc.add_paragraph()
        
        # 底部日期信息
        footer_config = self.cover_config.get('footer', {})
        if "completion_date" in doc_info:
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(doc_info["completion_date"])
            date_run.font.name = footer_config.get('font', '')
            date_run._element.rPr.rFonts.set(qn('w:eastAsia'), footer_config.get('font', ''))
            date_run._element.rPr.rFonts.set(qn('w:ascii'), footer_config.get('font', ''))
            date_run._element.rPr.rFonts.set(qn('w:hAnsi'), footer_config.get('font', ''))
            date_run.font.size = Pt(footer_config.get('size', 0))
            date_run.font.bold = footer_config.get('bold', False)
            date_para.alignment = get_alignment_constant(footer_config.get('alignment', ''))
        else:
            # 默认日期
            date_para = doc.add_paragraph()
            date_run = date_para.add_run("2026年4月20日")
            date_run.font.name = footer_config.get('font', '')
            date_run._element.rPr.rFonts.set(qn('w:eastAsia'), footer_config.get('font', ''))
            date_run._element.rPr.rFonts.set(qn('w:ascii'), footer_config.get('font', ''))
            date_run._element.rPr.rFonts.set(qn('w:hAnsi'), footer_config.get('font', ''))
            date_run.font.size = Pt(footer_config.get('size', 0))
            date_run.font.bold = footer_config.get('bold', True)
            date_para.alignment = get_alignment_constant(footer_config.get('alignment', ''))
        
        # 分页
        doc.add_page_break()
        ensure_consistent_page_setup(doc)
    
    def generate_english_cover(self, doc, doc_info):
        """生成英文封面页"""
        # 1. 提交声明
        submission_config = self.cover_config.get('submission_statement', {})
        if submission_config.get('text', ''):
            submission_para = doc.add_paragraph()
            submission_run = submission_para.add_run(submission_config['text'])
            submission_run.font.name = submission_config.get('font', '')
            submission_run._element.rPr.rFonts.set(qn('w:eastAsia'), submission_config.get('font', ''))
            submission_run._element.rPr.rFonts.set(qn('w:ascii'), submission_config.get('font', ''))
            submission_run._element.rPr.rFonts.set(qn('w:hAnsi'), submission_config.get('font', ''))
            submission_run.font.size = Pt(submission_config.get('size', 14))
            submission_run.font.bold = submission_config.get('bold', True)
            submission_para.alignment = get_alignment_constant(submission_config.get('alignment', ''))
            submission_para.paragraph_format.space_after = Pt(submission_config.get('space_after', 36))
        
        # 2. 论文标题
        title_config = self.cover_config.get('title', {})
        if title_config.get('text', ''):
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title_config['text'])
            title_run.font.name = title_config.get('font', '')
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_config.get('font', ''))
            title_run._element.rPr.rFonts.set(qn('w:ascii'), title_config.get('font', ''))
            title_run._element.rPr.rFonts.set(qn('w:hAnsi'), title_config.get('font', ''))
            title_run.font.size = Pt(title_config.get('size', 18))
            title_run.font.bold = title_config.get('bold', True)
            title_para.alignment = get_alignment_constant(title_config.get('alignment', ''))
            title_para.paragraph_format.space_after = Pt(title_config.get('space_after', 144))
        elif "title" in doc_info:
            # 兼容旧格式
            title_para = doc.add_paragraph()
            english_title = doc_info["title"].upper()
            title_run = title_para.add_run(english_title)
            title_run.font.name = self.cover_styles['title_english']['font']
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), self.cover_styles['title_english']['font'])
            title_run._element.rPr.rFonts.set(qn('w:ascii'), self.cover_styles['title_english']['font'])
            title_run._element.rPr.rFonts.set(qn('w:hAnsi'), self.cover_styles['title_english']['font'])
            title_run.font.size = Pt(self.cover_styles['title_english']['size'])
            title_run.font.bold = self.cover_styles['title_english']['bold']
            title_para.alignment = get_alignment_constant(self.cover_styles['title_english']['alignment'])
            title_para.paragraph_format.space_after = Pt(self.cover_styles['title_english']['space_after'])
        
        # 3. 论文信息
        info_items = self.cover_config.get('info_items', [])
        for item in info_items:
            label = item['label']
            field = item['field']
            content = doc_info.get(field, '')
            if content:
                info_para = doc.add_paragraph()
                label_run = info_para.add_run(f"{label}")
                label_run.font.name = self.cover_styles['info_label']['font']
                label_run._element.rPr.rFonts.set(qn('w:eastAsia'), self.cover_styles['info_label']['font'])
                label_run._element.rPr.rFonts.set(qn('w:ascii'), self.cover_styles['info_label']['font'])
                label_run._element.rPr.rFonts.set(qn('w:hAnsi'), self.cover_styles['info_label']['font'])
                label_run.font.size = Pt(self.cover_styles['info_label']['size'])
                label_run.font.bold = self.cover_styles['info_label']['bold']
                
                # 添加制表符和内容
                info_para.add_run("\t")
                content_run = info_para.add_run(content)
                content_run.font.name = self.cover_styles['info_content']['font']
                content_run._element.rPr.rFonts.set(qn('w:eastAsia'), self.cover_styles['info_content']['font'])
                content_run._element.rPr.rFonts.set(qn('w:ascii'), self.cover_styles['info_content']['font'])
                content_run._element.rPr.rFonts.set(qn('w:hAnsi'), self.cover_styles['info_content']['font'])
                content_run.font.size = Pt(self.cover_styles['info_content']['size'])
                
                # 设置段落格式
                info_para.paragraph_format.space_after = Pt(self.cover_styles['info_label']['space_after'])
                info_para.paragraph_format.line_spacing = self.cover_styles['info_label']['line_spacing']
                info_para.paragraph_format.left_indent = Pt(self.cover_styles['info_label'].get('left_indent', 0))
        
        # 4. 底部信息
        footer_config = self.cover_styles['footer']
        footer_text = footer_config.get('text', '')
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(footer_text)
        footer_run.font.name = footer_config['font']
        footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), footer_config['font'])
        footer_run._element.rPr.rFonts.set(qn('w:ascii'), footer_config['font'])
        footer_run._element.rPr.rFonts.set(qn('w:hAnsi'), footer_config['font'])
        footer_run.font.size = Pt(footer_config['size'])
        footer_run.font.bold = footer_config['bold']
        footer_para.alignment = get_alignment_constant(footer_config['alignment'])
        footer_para.paragraph_format.space_before = Pt(footer_config['space_before'])
        
        # 分页
        doc.add_page_break()
    
    def generate_originality_declaration(self, doc, doc_info):
        """生成原创性声明页"""
        # 1. 学校名称
        school_config = self.cover_config.get('school_name', {})
        if school_config.get('text', ''):
            school_para = doc.add_paragraph()
            school_run = school_para.add_run(school_config['text'])
            school_run.font.name = school_config.get('font', '')
            school_run._element.rPr.rFonts.set(qn('w:eastAsia'), school_config.get('font', ''))
            school_run._element.rPr.rFonts.set(qn('w:ascii'), school_config.get('font', ''))
            school_run._element.rPr.rFonts.set(qn('w:hAnsi'), school_config.get('font', ''))
            school_run.font.size = Pt(school_config.get('size', 16))
            school_run.font.bold = school_config.get('bold', True)
            school_para.alignment = get_alignment_constant(school_config.get('alignment', ''))
            school_para.paragraph_format.space_after = Pt(school_config.get('space_after', 12))
        
        # 2. 声明标题
        title_config = self.cover_config.get('title', {})
        if title_config.get('text', ''):
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title_config['text'])
            title_run.font.name = title_config.get('font', '')
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_config.get('font', ''))
            title_run._element.rPr.rFonts.set(qn('w:ascii'), title_config.get('font', ''))
            title_run._element.rPr.rFonts.set(qn('w:hAnsi'), title_config.get('font', ''))
            title_run.font.size = Pt(title_config.get('size', 16))
            title_run.font.bold = title_config.get('bold', True)
            title_para.alignment = get_alignment_constant(title_config.get('alignment', ''))
            title_para.paragraph_format.space_after = Pt(title_config.get('space_after', 24))
        
        # 3. 声明内容
        content_config = self.cover_config.get('content', {})
        if content_config.get('text', ''):
            content_para = doc.add_paragraph()
            content_run = content_para.add_run(content_config['text'])
            content_run.font.name = content_config.get('font', '')
            content_run._element.rPr.rFonts.set(qn('w:eastAsia'), content_config.get('font', ''))
            content_run._element.rPr.rFonts.set(qn('w:ascii'), content_config.get('font', ''))
            content_run._element.rPr.rFonts.set(qn('w:hAnsi'), content_config.get('font', ''))
            content_run.font.size = Pt(content_config.get('size', 14))
            content_run.font.bold = content_config.get('bold', False)
            content_para.alignment = get_alignment_constant(content_config.get('alignment', ''))
            content_para.paragraph_format.first_line_indent = Pt(content_config.get('first_line_indent', 0))
            content_para.paragraph_format.line_spacing = content_config.get('line_spacing', 1.5)
            content_para.paragraph_format.space_after = Pt(content_config.get('space_after', 36))
        
        # 4. 签名区域
        signature_config = self.cover_config.get('signature_section', {})
        if signature_config.get('text', ''):
            signature_para = doc.add_paragraph()
            signature_run = signature_para.add_run(signature_config['text'])
            signature_run.font.name = signature_config.get('font', '')
            signature_run._element.rPr.rFonts.set(qn('w:eastAsia'), signature_config.get('font', ''))
            signature_run._element.rPr.rFonts.set(qn('w:ascii'), signature_config.get('font', ''))
            signature_run._element.rPr.rFonts.set(qn('w:hAnsi'), signature_config.get('font', ''))
            signature_run.font.size = Pt(signature_config.get('size', 14))
            signature_run.font.bold = signature_config.get('bold', True)
            signature_para.alignment = get_alignment_constant(signature_config.get('alignment', ''))
            signature_para.paragraph_format.space_after = Pt(signature_config.get('space_after', 12))
        
        # 5. 日期区域
        date_config = self.cover_config.get('date_section', {})
        if date_config.get('text', ''):
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(date_config['text'])
            date_run.font.name = date_config.get('font', '宋体')
            date_run._element.rPr.rFonts.set(qn('w:eastAsia'), date_config.get('font', '宋体'))
            date_run._element.rPr.rFonts.set(qn('w:ascii'), date_config.get('font', '宋体'))
            date_run._element.rPr.rFonts.set(qn('w:hAnsi'), date_config.get('font', '宋体'))
            date_run.font.size = Pt(date_config.get('size', 14))
            date_run.font.bold = date_config.get('bold', True)
            date_para.alignment = get_alignment_constant(date_config.get('alignment', 'right'))
            date_para.paragraph_format.space_after = Pt(date_config.get('space_after', 48))
        
        # 6. 使用授权书标题
        auth_title_config = self.cover_config.get('authorization_title', {})
        if auth_title_config.get('text', ''):
            auth_title_para = doc.add_paragraph()
            auth_title_run = auth_title_para.add_run(auth_title_config['text'])
            auth_title_run.font.name = auth_title_config.get('font', '黑体')
            auth_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), auth_title_config.get('font', '黑体'))
            auth_title_run._element.rPr.rFonts.set(qn('w:ascii'), auth_title_config.get('font', '黑体'))
            auth_title_run._element.rPr.rFonts.set(qn('w:hAnsi'), auth_title_config.get('font', '黑体'))
            auth_title_run.font.size = Pt(auth_title_config.get('size', 16))
            auth_title_run.font.bold = auth_title_config.get('bold', True)
            auth_title_para.alignment = get_alignment_constant(auth_title_config.get('alignment', 'center'))
            auth_title_para.paragraph_format.space_after = Pt(auth_title_config.get('space_after', 24))
        
        # 7. 授权书内容
        auth_content_config = self.cover_config.get('authorization_content', {})
        if auth_content_config.get('text', ''):
            auth_content_para = doc.add_paragraph()
            auth_content_run = auth_content_para.add_run(auth_content_config['text'])
            auth_content_run.font.name = auth_content_config.get('font', '宋体')
            auth_content_run._element.rPr.rFonts.set(qn('w:eastAsia'), auth_content_config.get('font', '宋体'))
            auth_content_run._element.rPr.rFonts.set(qn('w:ascii'), auth_content_config.get('font', '宋体'))
            auth_content_run._element.rPr.rFonts.set(qn('w:hAnsi'), auth_content_config.get('font', '宋体'))
            auth_content_run.font.size = Pt(auth_content_config.get('size', 14))
            auth_content_run.font.bold = auth_content_config.get('bold', False)
            auth_content_para.alignment = get_alignment_constant(auth_content_config.get('alignment', 'left'))
            auth_content_para.paragraph_format.first_line_indent = Pt(auth_content_config.get('first_line_indent', 0))
            auth_content_para.paragraph_format.line_spacing = auth_content_config.get('line_spacing', 1.5)
            auth_content_para.paragraph_format.space_after = Pt(auth_content_config.get('space_after', 24))
        
        # 8. 保密选项
        confidentiality_config = self.cover_config.get('confidentiality_options', {})
        if confidentiality_config.get('text', ''):
            confidentiality_para = doc.add_paragraph()
            
            # 拆分文本并为论文类型添加加粗
            text = confidentiality_config['text']
            parts = text.split('□')
            
            for i, part in enumerate(parts):
                if i == 0:
                    # 第一部分：“本学位论文属于 ：”
                    run = confidentiality_para.add_run(part)
                    run.font.name = confidentiality_config.get('font', '宋体')
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), confidentiality_config.get('font', '宋体'))
                    run._element.rPr.rFonts.set(qn('w:ascii'), confidentiality_config.get('font', '宋体'))
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), confidentiality_config.get('font', '宋体'))
                    run.font.size = Pt(confidentiality_config.get('size', 14))
                    run.font.bold = confidentiality_config.get('bold', False)
                else:
                    # 添加方框 - 使用更大的字体且明确不加粗
                    square_run = confidentiality_para.add_run('□')
                    square_run.font.name = confidentiality_config.get('font', '宋体')
                    square_run._element.rPr.rFonts.set(qn('w:eastAsia'), confidentiality_config.get('font', '宋体'))
                    square_run._element.rPr.rFonts.set(qn('w:ascii'), confidentiality_config.get('font', '宋体'))
                    square_run._element.rPr.rFonts.set(qn('w:hAnsi'), confidentiality_config.get('font', '宋体'))
                    square_run.font.size = Pt(16)  # 增大方框字体
                    square_run.font.bold = False  # 明确设置方框不加粗
                    # 确保其他样式也正常
                    square_run.font.italic = False
                    square_run.font.underline = False
                    
                    # 处理论文类型名称 - 只加粗类型名称，其余保持正常
                    if part.startswith('公开论文'):
                        # 只加粗"公开论文"
                        run = confidentiality_para.add_run('公开论文')
                        run.font.name = confidentiality_config.get('font', '宋体')
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), confidentiality_config.get('font', '宋体'))
                        run._element.rPr.rFonts.set(qn('w:ascii'), confidentiality_config.get('font', '宋体'))
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), confidentiality_config.get('font', '宋体'))
                        run.font.size = Pt(confidentiality_config.get('size', 14))
                        run.font.bold = True
                        # 剩余部分正常
                        if len(part) > 4:
                            rest = part[4:]
                            run = confidentiality_para.add_run(rest)
                            run.font.name = confidentiality_config.get('font', '宋体')
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), confidentiality_config.get('font', '宋体'))
                            run._element.rPr.rFonts.set(qn('w:ascii'), confidentiality_config.get('font', '宋体'))
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), confidentiality_config.get('font', '宋体'))
                            run.font.size = Pt(confidentiality_config.get('size', 14))
                            run.font.bold = False
                    elif part.startswith('内部论文'):
                        # 只加粗"内部论文"
                        run = confidentiality_para.add_run('内部论文')
                        run.font.name = confidentiality_config.get('font', '宋体')
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), confidentiality_config.get('font', '宋体'))
                        run._element.rPr.rFonts.set(qn('w:ascii'), confidentiality_config.get('font', '宋体'))
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), confidentiality_config.get('font', '宋体'))
                        run.font.size = Pt(confidentiality_config.get('size', 14))
                        run.font.bold = True
                        # 剩余部分正常
                        if len(part) > 4:
                            rest = part[4:]
                            run = confidentiality_para.add_run(rest)
                            run.font.name = confidentiality_config.get('font', '宋体')
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), confidentiality_config.get('font', '宋体'))
                            run._element.rPr.rFonts.set(qn('w:ascii'), confidentiality_config.get('font', '宋体'))
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), confidentiality_config.get('font', '宋体'))
                            run.font.size = Pt(confidentiality_config.get('size', 14))
                            run.font.bold = False
                    elif part.startswith('秘密论文'):
                        # 只加粗"秘密论文"
                        run = confidentiality_para.add_run('秘密论文')
                        run.font.name = confidentiality_config.get('font', '宋体')
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), confidentiality_config.get('font', '宋体'))
                        run._element.rPr.rFonts.set(qn('w:ascii'), confidentiality_config.get('font', '宋体'))
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), confidentiality_config.get('font', '宋体'))
                        run.font.size = Pt(confidentiality_config.get('size', 14))
                        run.font.bold = True
                        # 剩余部分正常
                        if len(part) > 4:
                            rest = part[4:]
                            run = confidentiality_para.add_run(rest)
                            run.font.name = confidentiality_config.get('font', '宋体')
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), confidentiality_config.get('font', '宋体'))
                            run._element.rPr.rFonts.set(qn('w:ascii'), confidentiality_config.get('font', '宋体'))
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), confidentiality_config.get('font', '宋体'))
                            run.font.size = Pt(confidentiality_config.get('size', 14))
                            run.font.bold = False
                    elif part.startswith('机密论文'):
                        # 只加粗"机密论文"
                        run = confidentiality_para.add_run('机密论文')
                        run.font.name = confidentiality_config.get('font', '宋体')
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), confidentiality_config.get('font', '宋体'))
                        run._element.rPr.rFonts.set(qn('w:ascii'), confidentiality_config.get('font', '宋体'))
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), confidentiality_config.get('font', '宋体'))
                        run.font.size = Pt(confidentiality_config.get('size', 14))
                        run.font.bold = True
                        # 剩余部分正常
                        if len(part) > 4:
                            rest = part[4:]
                            run = confidentiality_para.add_run(rest)
                            run.font.name = confidentiality_config.get('font', '宋体')
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), confidentiality_config.get('font', '宋体'))
                            run._element.rPr.rFonts.set(qn('w:ascii'), confidentiality_config.get('font', '宋体'))
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), confidentiality_config.get('font', '宋体'))
                            run.font.size = Pt(confidentiality_config.get('size', 14))
                            run.font.bold = False
                    else:
                        # 其他情况，正常显示
                        run = confidentiality_para.add_run(part)
                        run.font.name = confidentiality_config.get('font', '宋体')
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), confidentiality_config.get('font', '宋体'))
                        run._element.rPr.rFonts.set(qn('w:ascii'), confidentiality_config.get('font', '宋体'))
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), confidentiality_config.get('font', '宋体'))
                        run.font.size = Pt(confidentiality_config.get('size', 14))
                        run.font.bold = False
            
            confidentiality_para.alignment = get_alignment_constant(confidentiality_config.get('alignment', 'left'))
            confidentiality_para.paragraph_format.line_spacing = confidentiality_config.get('line_spacing', 1.5)
            confidentiality_para.paragraph_format.space_after = Pt(confidentiality_config.get('space_after', 36))
        
        # 9. 授权书签名区域
        auth_signature_config = self.cover_config.get('authorization_signature', {})
        if auth_signature_config.get('text', ''):
            auth_signature_para = doc.add_paragraph()
            auth_signature_run = auth_signature_para.add_run(auth_signature_config['text'])
            auth_signature_run.font.name = auth_signature_config.get('font', '宋体')
            auth_signature_run._element.rPr.rFonts.set(qn('w:eastAsia'), auth_signature_config.get('font', '宋体'))
            auth_signature_run._element.rPr.rFonts.set(qn('w:ascii'), auth_signature_config.get('font', '宋体'))
            auth_signature_run._element.rPr.rFonts.set(qn('w:hAnsi'), auth_signature_config.get('font', '宋体'))
            auth_signature_run.font.size = Pt(auth_signature_config.get('size', 14))
            auth_signature_run.font.bold = auth_signature_config.get('bold', True)
            auth_signature_para.alignment = get_alignment_constant(auth_signature_config.get('alignment', 'left'))
            auth_signature_para.paragraph_format.space_after = Pt(auth_signature_config.get('space_after', 12))
        
        # 10. 授权书日期区域
        auth_date_config = self.cover_config.get('authorization_date', {})
        if auth_date_config.get('text', ''):
            auth_date_para = doc.add_paragraph()
            auth_date_run = auth_date_para.add_run(auth_date_config['text'])
            auth_date_run.font.name = auth_date_config.get('font', '宋体')
            auth_date_run._element.rPr.rFonts.set(qn('w:eastAsia'), auth_date_config.get('font', '宋体'))
            auth_date_run._element.rPr.rFonts.set(qn('w:ascii'), auth_date_config.get('font', '宋体'))
            auth_date_run._element.rPr.rFonts.set(qn('w:hAnsi'), auth_date_config.get('font', '宋体'))
            auth_date_run.font.size = Pt(auth_date_config.get('size', 14))
            auth_date_run.font.bold = auth_date_config.get('bold', True)
            auth_date_para.alignment = get_alignment_constant(auth_date_config.get('alignment', 'left'))
            auth_date_para.paragraph_format.space_after = Pt(auth_date_config.get('space_after', 24))
        
        # 分页
        doc.add_page_break()


class AbstractGenerator:
    """摘要生成器"""
    
    def __init__(self, config):
        # 从配置文件加载样式
        self.abstract_styles = config.get('styles', {})
        # 优先使用abstract_config中的样式
        if 'abstract' in config and 'styles' in config['abstract']:
            self.abstract_styles = config['abstract']['styles']
        # 如果styles是列表，尝试找到抽象样式
        if isinstance(self.abstract_styles, list):
            for style in self.abstract_styles:
                if isinstance(style, dict) and 'chinese_title' in style:
                    self.abstract_styles = style
                    break
            # 如果遍历后仍然是列表，转换为空字典
            if isinstance(self.abstract_styles, list):
                self.abstract_styles = {}
        # 对齐方式映射
        self.align_mapping = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
    
    def generate_chinese_abstract(self, doc, abstract_info):
        """生成中文摘要"""
        # 获取中文摘要信息
        chinese_info = abstract_info.get('chinese', {})
        # 摘要标题（两字间空一格）
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(chinese_info.get('title', '摘  要'))
        title_style = self.abstract_styles.get('chinese_title', {})
        title_run.font.name = title_style.get('font', '黑体')
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_style.get('font', '黑体'))
        title_run._element.rPr.rFonts.set(qn('w:ascii'), title_style.get('font', '黑体'))
        title_run._element.rPr.rFonts.set(qn('w:hAnsi'), title_style.get('font', '黑体'))
        title_run.font.size = Pt(title_style.get('size', 16))
        title_run.font.bold = title_style.get('bold', True)
        title_para.alignment = self.align_mapping.get(title_style.get('alignment', 'center'), WD_ALIGN_PARAGRAPH.CENTER)
        title_para.paragraph_format.space_before = Pt(title_style.get('space_before', 24))
        title_para.paragraph_format.space_after = Pt(title_style.get('space_after', 18))
        title_para.paragraph_format.line_spacing = title_style.get('line_spacing', 1.0)
        
        # 摘要正文
        if "content" in chinese_info:
            # 将内容按换行符分割成多个段落
            paragraphs = chinese_info["content"].split('\n')
            content_style = self.abstract_styles.get('chinese_content', {})
            for i, para_text in enumerate(paragraphs):
                if not para_text.strip():  # 跳过空行
                    continue
                content_para = doc.add_paragraph()
                content_run = content_para.add_run(para_text)
                content_run.font.name = content_style.get('font', '宋体')
                content_run._element.rPr.rFonts.set(qn('w:eastAsia'), content_style.get('font', '宋体'))
                content_run._element.rPr.rFonts.set(qn('w:ascii'), content_style.get('font', '宋体'))
                content_run._element.rPr.rFonts.set(qn('w:hAnsi'), content_style.get('font', '宋体'))
                content_run.font.size = Pt(content_style.get('size', 12))
                content_run.font.bold = content_style.get('bold', False)
                content_para.alignment = self.align_mapping.get(content_style.get('alignment', 'left'), WD_ALIGN_PARAGRAPH.LEFT)
                content_para.paragraph_format.first_line_indent = Pt(content_style.get('first_line_indent', 24))
                content_para.paragraph_format.space_before = Pt(content_style.get('space_before', 0))
                content_para.paragraph_format.space_after = Pt(content_style.get('space_after', 0))
                content_para.paragraph_format.line_spacing = Pt(content_style.get('line_spacing', 20))
        
        # 关键词标签（前添加空行）
        doc.add_paragraph()  # 添加空行
        keywords_label_para = doc.add_paragraph()
        keywords_label_run = keywords_label_para.add_run("关键词：")
        keywords_label_style = self.abstract_styles.get('chinese_keywords_label', {})
        keywords_label_run.font.name = keywords_label_style.get('font', '宋体')
        keywords_label_run._element.rPr.rFonts.set(qn('w:eastAsia'), keywords_label_style.get('font', '宋体'))
        keywords_label_run._element.rPr.rFonts.set(qn('w:ascii'), keywords_label_style.get('font', '宋体'))
        keywords_label_run._element.rPr.rFonts.set(qn('w:hAnsi'), keywords_label_style.get('font', '宋体'))
        keywords_label_run.font.size = Pt(keywords_label_style.get('size', 12))
        keywords_label_run.font.bold = True  # 强制加粗
        keywords_label_para.alignment = self.align_mapping.get(keywords_label_style.get('alignment', 'left'), WD_ALIGN_PARAGRAPH.LEFT)
        
        # 关键词内容
        if "keywords" in chinese_info:
            keywords_text = "，".join(chinese_info["keywords"])
            keywords_content_run = keywords_label_para.add_run(keywords_text)
            keywords_content_style = self.abstract_styles.get('chinese_keywords_content', {})
            keywords_content_run.font.name = keywords_content_style.get('font', '宋体')
            keywords_content_run._element.rPr.rFonts.set(qn('w:eastAsia'), keywords_content_style.get('font', '宋体'))
            keywords_content_run._element.rPr.rFonts.set(qn('w:ascii'), keywords_content_style.get('font', '宋体'))
            keywords_content_run._element.rPr.rFonts.set(qn('w:hAnsi'), keywords_content_style.get('font', '宋体'))
            keywords_content_run.font.size = Pt(keywords_content_style.get('size', 12))
            keywords_content_run.font.bold = keywords_content_style.get('bold', False)
        
        # 分页
        doc.add_page_break()
        ensure_consistent_page_setup(doc)
    
    def generate_english_abstract(self, doc, abstract_info):
        """生成英文摘要"""
        # 获取英文摘要信息
        english_info = abstract_info.get('english', {})
        # 摘要标题
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(english_info.get('title', ''))
        title_style = self.abstract_styles.get('english_title', {})
        title_run.font.name = title_style.get('font', 'Arial')
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_style.get('font', 'Arial'))
        title_run._element.rPr.rFonts.set(qn('w:ascii'), title_style.get('font', 'Arial'))
        title_run._element.rPr.rFonts.set(qn('w:hAnsi'), title_style.get('font', 'Arial'))
        title_run.font.size = Pt(title_style.get('size', 16))
        title_run.font.bold = title_style.get('bold', False)
        title_para.alignment = self.align_mapping.get(title_style.get('alignment', 'center'), WD_ALIGN_PARAGRAPH.CENTER)
        title_para.paragraph_format.space_before = Pt(title_style.get('space_before', 24))
        title_para.paragraph_format.space_after = Pt(title_style.get('space_after', 18))
        title_para.paragraph_format.line_spacing = title_style.get('line_spacing', 1.0)
        
        # 摘要正文
        if "content" in english_info:
            # 将内容按换行符分割成多个段落
            paragraphs = english_info["content"].split('\n')
            content_style = self.abstract_styles.get('english_content', {})
            for i, para_text in enumerate(paragraphs):
                if not para_text.strip():  # 跳过空行
                    continue
                content_para = doc.add_paragraph()
                content_run = content_para.add_run(para_text)
                content_run.font.name = content_style.get('font', 'Times New Roman')
                content_run._element.rPr.rFonts.set(qn('w:eastAsia'), content_style.get('font', 'Times New Roman'))
                content_run._element.rPr.rFonts.set(qn('w:ascii'), content_style.get('font', 'Times New Roman'))
                content_run._element.rPr.rFonts.set(qn('w:hAnsi'), content_style.get('font', 'Times New Roman'))
                content_run.font.size = Pt(content_style.get('size', 12))
                content_run.font.bold = content_style.get('bold', False)
                content_para.alignment = self.align_mapping.get(content_style.get('alignment', 'left'), WD_ALIGN_PARAGRAPH.LEFT)
                content_para.paragraph_format.first_line_indent = Pt(content_style.get('first_line_indent', 24))
                content_para.paragraph_format.space_before = Pt(content_style.get('space_before', 0))
                content_para.paragraph_format.space_after = Pt(content_style.get('space_after', 0))
                content_para.paragraph_format.line_spacing = Pt(content_style.get('line_spacing', 18))
        
        # 关键词标签（前添加空行）
        doc.add_paragraph()  # 添加空行
        keywords_label_para = doc.add_paragraph()
        keywords_label_run = keywords_label_para.add_run("Key words: ")
        keywords_label_style = self.abstract_styles.get('english_keywords_label', {})
        keywords_label_run.font.name = keywords_label_style.get('font', 'Times New Roman')
        keywords_label_run._element.rPr.rFonts.set(qn('w:eastAsia'), keywords_label_style.get('font', 'Times New Roman'))
        keywords_label_run._element.rPr.rFonts.set(qn('w:ascii'), keywords_label_style.get('font', 'Times New Roman'))
        keywords_label_run._element.rPr.rFonts.set(qn('w:hAnsi'), keywords_label_style.get('font', 'Times New Roman'))
        keywords_label_run.font.size = Pt(keywords_label_style.get('size', 12))
        keywords_label_run.font.bold = True  # 强制加粗
        keywords_label_para.alignment = self.align_mapping.get(keywords_label_style.get('alignment', 'left'), WD_ALIGN_PARAGRAPH.LEFT)
        
        # 关键词内容
        if "keywords" in english_info:
            keywords_text = ", ".join(english_info["keywords"])
            keywords_content_run = keywords_label_para.add_run(keywords_text)
            keywords_content_style = self.abstract_styles.get('english_keywords_content', {})
            keywords_content_run.font.name = keywords_content_style.get('font', 'Times New Roman')
            keywords_content_run._element.rPr.rFonts.set(qn('w:eastAsia'), keywords_content_style.get('font', 'Times New Roman'))
            keywords_content_run._element.rPr.rFonts.set(qn('w:ascii'), keywords_content_style.get('font', 'Times New Roman'))
            keywords_content_run._element.rPr.rFonts.set(qn('w:hAnsi'), keywords_content_style.get('font', 'Times New Roman'))
            keywords_content_run.font.size = Pt(keywords_content_style.get('size', 12))
            keywords_content_run.font.bold = keywords_content_style.get('bold', False)


def create_custom_styles(doc, config=None):
    """创建文档所需的自定义样式"""
    # 从配置文件获取样式设置
    text_styles = config.get('text_styles', {}) if config else {}
    image_config = config.get('image', {}) if config else {}
    fig_config = image_config.get('figure', {})
    
    styles = doc.styles
    
    # 创建正文样式
    body_config = text_styles.get('body', {})
    body_style = styles.add_style('\u200B正文', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = body_config.get('font', '宋体')
    body_style._element.rPr.rFonts.set(qn('w:eastAsia'), body_config.get('font', '宋体'))
    body_style._element.rPr.rFonts.set(qn('w:ascii'), body_config.get('english_font', 'Times New Roman'))
    body_style._element.rPr.rFonts.set(qn('w:hAnsi'), body_config.get('english_font', 'Times New Roman'))
    body_style.font.size = Pt(body_config.get('size', 12))
    body_style.font.bold = body_config.get('bold', False)
    body_style.paragraph_format.first_line_indent = Pt(body_config.get('first_line_indent', 24))
    body_style.paragraph_format.line_spacing = body_config.get('line_spacing', 1.5)
    body_style.paragraph_format.space_before = Pt(body_config.get('space_before', 0))
    body_style.paragraph_format.space_after = Pt(body_config.get('space_after', 0))
    
    # 创建一级标题样式
    h1_config = text_styles.get('heading1', {})
    h1_style = styles.add_style('\u200D一级标题', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = h1_config.get('font', '黑体')
    h1_style._element.rPr.rFonts.set(qn('w:eastAsia'), h1_config.get('font', '黑体'))
    h1_style._element.rPr.rFonts.set(qn('w:ascii'), h1_config.get('english_font', 'Arial'))
    h1_style._element.rPr.rFonts.set(qn('w:hAnsi'), h1_config.get('english_font', 'Arial'))
    h1_style.font.size = Pt(h1_config.get('size', 15))
    h1_style.font.bold = h1_config.get('bold', True)
    h1_style.paragraph_format.first_line_indent = Pt(h1_config.get('first_line_indent', 0))
    h1_style.paragraph_format.line_spacing = h1_config.get('line_spacing', 1.5)
    h1_style.paragraph_format.space_before = Pt(h1_config.get('space_before', 24))
    h1_style.paragraph_format.space_after = Pt(h1_config.get('space_after', 18))
    
    # 创建二级标题样式
    h2_config = text_styles.get('heading2', {})
    h2_style = styles.add_style('\u200e二级标题', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = h2_config.get('font', '黑体')
    h2_style._element.rPr.rFonts.set(qn('w:eastAsia'), h2_config.get('font', '黑体'))
    h2_style._element.rPr.rFonts.set(qn('w:ascii'), h2_config.get('english_font', 'Arial'))
    h2_style._element.rPr.rFonts.set(qn('w:hAnsi'), h2_config.get('english_font', 'Arial'))
    h2_style.font.size = Pt(h2_config.get('size', 14))
    h2_style.font.bold = h2_config.get('bold', True)
    h2_style.paragraph_format.first_line_indent = Pt(h2_config.get('first_line_indent', 0))
    h2_style.paragraph_format.line_spacing = h2_config.get('line_spacing', 1.5)
    h2_style.paragraph_format.space_before = Pt(h2_config.get('space_before', 18))
    h2_style.paragraph_format.space_after = Pt(h2_config.get('space_after', 12))
    
    # 创建三级标题样式
    h3_config = text_styles.get('heading3', {})
    h3_style = styles.add_style('\u200F三级标题', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.name = h3_config.get('font', '宋体')
    h3_style._element.rPr.rFonts.set(qn('w:eastAsia'), h3_config.get('font', '宋体'))
    h3_style._element.rPr.rFonts.set(qn('w:ascii'), h3_config.get('english_font', 'Times New Roman'))
    h3_style._element.rPr.rFonts.set(qn('w:hAnsi'), h3_config.get('english_font', 'Times New Roman'))
    h3_style.font.size = Pt(h3_config.get('size', 13))
    h3_style.font.bold = h3_config.get('bold', True)
    h3_style.paragraph_format.first_line_indent = Pt(h3_config.get('first_line_indent', 21.6))
    h3_style.paragraph_format.line_spacing = h3_config.get('line_spacing', 1.5)
    h3_style.paragraph_format.space_before = Pt(h3_config.get('space_before', 12))
    h3_style.paragraph_format.space_after = Pt(h3_config.get('space_after', 6))
    
    # 创建图题样式
    fig_style = styles.add_style('\u200C图题', WD_STYLE_TYPE.PARAGRAPH)
    fig_style.font.name = fig_config.get('font', '宋体')
    fig_style._element.rPr.rFonts.set(qn('w:eastAsia'), fig_config.get('font', '宋体'))
    fig_style._element.rPr.rFonts.set(qn('w:ascii'), fig_config.get('english_font', 'Times New Roman'))
    fig_style._element.rPr.rFonts.set(qn('w:hAnsi'), fig_config.get('english_font', 'Times New Roman'))
    fig_style.font.size = Pt(fig_config.get('size', 10.5))
    fig_style.font.bold = fig_config.get('bold', True)
    fig_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_style.paragraph_format.first_line_indent = Pt(0)
    fig_style.paragraph_format.line_spacing = fig_config.get('line_spacing', 1.5)
    fig_style.paragraph_format.space_before = Pt(fig_config.get('space_before', 0))
    fig_style.paragraph_format.space_after = Pt(fig_config.get('space_after', 6))
    
    # 创建参考文献样式
    ref_config = text_styles.get('reference', {})
    ref_style = styles.add_style('\u206F参考文献', WD_STYLE_TYPE.PARAGRAPH)
    ref_style.font.name = ref_config.get('font', '宋体')
    ref_style._element.rPr.rFonts.set(qn('w:eastAsia'), ref_config.get('font', '宋体'))
    ref_style._element.rPr.rFonts.set(qn('w:ascii'), ref_config.get('english_font', 'Times New Roman'))
    ref_style._element.rPr.rFonts.set(qn('w:hAnsi'), ref_config.get('english_font', 'Times New Roman'))
    ref_style.font.size = Pt(ref_config.get('size', 10.5))
    ref_style.font.bold = ref_config.get('bold', False)
    ref_style.paragraph_format.first_line_indent = Pt(ref_config.get('first_line_indent', 21.6))
    ref_style.paragraph_format.line_spacing = ref_config.get('line_spacing', 1.5)
    ref_style.paragraph_format.hanging_indent = Pt(ref_config.get('hanging_indent', 21.6))
    ref_style.paragraph_format.space_before = Pt(ref_config.get('space_before', 0))
    ref_style.paragraph_format.space_after = Pt(ref_config.get('space_after', 0))
    
    # 创建参考文献标题样式
    ref_title_config = text_styles.get('reference_title', {})
    ref_title_style = styles.add_style('\u2011参考文献标题', WD_STYLE_TYPE.PARAGRAPH)
    ref_title_style.font.name = ref_title_config.get('font', '黑体')
    ref_title_style._element.rPr.rFonts.set(qn('w:eastAsia'), ref_title_config.get('font', '黑体'))
    ref_title_style._element.rPr.rFonts.set(qn('w:ascii'), ref_title_config.get('english_font', 'Times New Roman'))
    ref_title_style._element.rPr.rFonts.set(qn('w:hAnsi'), ref_title_config.get('english_font', 'Times New Roman'))
    ref_title_style.font.size = Pt(ref_title_config.get('size', 16))
    ref_title_style.font.bold = ref_title_config.get('bold', True)
    ref_title_style.paragraph_format.first_line_indent = Pt(ref_title_config.get('first_line_indent', 0))
    ref_title_style.paragraph_format.line_spacing = ref_title_config.get('line_spacing', 1.5)
    ref_title_style.paragraph_format.space_before = Pt(ref_title_config.get('space_before', 24))
    ref_title_style.paragraph_format.space_after = Pt(ref_title_config.get('space_after', 18))
    
    # 创建全文大标题样式
    global_title_config = text_styles.get('global_title', {})
    global_title_style = styles.add_style('\u2010全文大标题', WD_STYLE_TYPE.PARAGRAPH)
    global_title_style.font.name = global_title_config.get('font', '黑体')
    global_title_style._element.rPr.rFonts.set(qn('w:eastAsia'), global_title_config.get('font', '黑体'))
    global_title_style._element.rPr.rFonts.set(qn('w:ascii'), global_title_config.get('english_font', 'Times New Roman'))
    global_title_style._element.rPr.rFonts.set(qn('w:hAnsi'), global_title_config.get('english_font', 'Times New Roman'))
    global_title_style.font.size = Pt(global_title_config.get('size', 22))
    global_title_style.font.bold = global_title_config.get('bold', True)
    global_title_style.paragraph_format.first_line_indent = Pt(global_title_config.get('first_line_indent', 0))
    global_title_style.paragraph_format.line_spacing = global_title_config.get('line_spacing', 1.5)
    global_title_style.paragraph_format.space_before = Pt(global_title_config.get('space_before', 24))
    global_title_style.paragraph_format.space_after = Pt(global_title_config.get('space_after', 18))
    # 设置对齐方式
    alignment = global_title_config.get('alignment', '')
    if alignment == 'center':
        global_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_word_template(config_file=None, output_filename=None):
    """创建Word模板"""
    global image_counter, small_image_count, large_image_count
    # 重置计数器和统计变量
    image_counter = 1
    small_image_count = 0
    large_image_count = 0
    
    # 创建标题收集器
    title_collector = TitleCollector()
    
    # 获取当前脚本的目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    if config_file is None:
        # 使用CONFIG中的默认配置
        config_file = CONFIG['default_config_file']
    
    # 读取配置文件
    try:
        config = load_multi_config_file(config_file)
        # 如果输出文件名未指定，从配置文件中读取
        if output_filename is None:
            if "generation_config" in config and "output_filename" in config["generation_config"]:
                config_output_filename = config["generation_config"]["output_filename"]
                # 如果输出文件名不是绝对路径，转换为相对于脚本文件的绝对路径
                if not os.path.isabs(config_output_filename):
                    output_filename = os.path.join(script_dir, config_output_filename)
                else:
                    output_filename = config_output_filename
            else:
                # 使用默认输出文件名
                output_filename = os.path.join(script_dir, CONFIG['default_output_filename'])
    except FileNotFoundError:
        print(f"配置文件 {config_file} 未找到")
        return None
    except Exception as e:
        print(f"读取配置文件时出错: {e}")
        return None
    
    doc = Document()
    
    # 创建自定义样式
    create_custom_styles(doc, config)
    
    # 设置A4页面大小和标准边距
    ensure_consistent_page_setup(doc, config)
    
    # 直接添加内容结构（跳过封面、摘要、目录等）
    if "content_structure" in config:
        content = config["content_structure"]
        
        # 添加章节
        if "sections" in content:
            for i, section in enumerate(content["sections"]):
                # 处理图片类型的section
                if section["type"] == "图片":
                    if "paths" in section and "captions" in section:
                        image_paths = section.get("paths", [])
                        captions = section.get("captions", [])
                        if image_paths:
                            add_adaptive_images(doc, image_paths, captions, config)
                    else:
                        image_path = section.get("path", "")
                        caption = section.get("caption", "")
                        if image_path:
                            add_adaptive_images(doc, [image_path], [caption], config)
                    continue
                
                # 处理表格类型的section
                elif section["type"] == "表格":
                    table_name = section.get("table_name", "")
                    headers = section.get("headers", [])
                    data_rows = section.get("data_rows", [])
                    caption = section.get("caption", "")
                    if headers and data_rows:
                        table_data = [headers] + data_rows
                        create_three_line_table(doc, table_data, caption, config)
                        # 添加空行，与后续内容保持间距
                        doc.add_paragraph()
                    continue
                
                # 处理普通文本类型的section
                if "text" not in section:
                    continue
                
                full_style_name = STYLE_MAPPING.get(section["type"], f"\u200B{section['type']}")
                alignment_str = section.get("alignment", None)
                alignment = get_alignment_constant(alignment_str) if alignment_str else None
                add_paragraph_with_style(doc, section["text"], full_style_name, alignment=alignment)
                
                # 处理列表格式
                if "list" in section:
                    list_data = section["list"]
                    if "items" in list_data:
                        level = list_data.get("level", 1)
                        add_numbered_list(doc, list_data["items"], level)
                
                # 处理普通段落格式
                elif "content_paragraphs" in section:
                    for para_data in section["content_paragraphs"]:
                        style_name = para_data.get("style", "正文")
                        full_style_name = STYLE_MAPPING.get(style_name, f"\u200B{style_name}")
                        alignment_str = para_data.get("alignment", None)
                        alignment = get_alignment_constant(alignment_str) if alignment_str else None
                        add_paragraph_with_style(doc, para_data["text"], full_style_name, alignment=alignment)
        
        # 添加示例内容
        if "examples" in content:
            examples = content["examples"]
            
            # 添加图题示例
            if "figure_caption" in examples:
                fig_cap = examples["figure_caption"]
                add_paragraph_with_style(doc, fig_cap["text"], f"\u200C{fig_cap['style']}")
            
            # 添加参考文献示例
            if "references_section" in examples:
                ref_section = examples["references_section"]
                add_paragraph_with_style(doc, ref_section["title"], f"\u200D{ref_section['title_style']}")
                for ref_item in ref_section["items"]:
                    add_paragraph_with_style(doc, ref_item["text"], f"\u206F{ref_item['style']}")
    
    add_page_numbers(doc, None, config)
    return doc, output_filename



if __name__ == "__main__":
    import sys
    
    # 支持命令行参数：配置文件名和输出文件名
    config_file = None
    output_filename = None
    
    if len(sys.argv) >= 2:
        config_file = sys.argv[1]
    if len(sys.argv) >= 3:
        output_filename = sys.argv[2]
    
    # 如果没有提供参数，使用默认值
    if config_file is None:
        # 使用CONFIG中的默认配置
        config_file = CONFIG['default_config_file']
    
    try:
        result = create_word_template(config_file, output_filename)
        if result:
            doc, final_filename = result
            output_filename = final_filename
            
            # 计算参考文献数量并检测重复
            references_count = 0
            duplicate_references = []
            try:
                config = load_multi_config_file(config_file if config_file else CONFIG['default_config_file'])
                if "_reference_manager" in config:
                    ref_manager = config["_reference_manager"]
                    all_references = ref_manager.get_all_references()
                    
                    # 检测重复参考文献
                    unique_references = []
                    ref_texts = []
                    duplicate_references = []
                    original_refs = []
                    
                    for number, ref_text in all_references:
                        original_refs.append((number, ref_text))
                        if ref_text not in ref_texts:
                            unique_references.append((number, ref_text))
                            ref_texts.append(ref_text)
                        else:
                            duplicate_references.append((number, ref_text))
                    
                    # 输出全部参考文献（移除多余序号，只保留参考文献自身编号）
                    # 参考文献列表已在引用管理器中处理
                    
                    if unique_references:
                        # 使用配置的参考文献标题样式添加章节标题
                        ref_title_style_name = '\u2011参考文献标题'
                        para = doc.add_paragraph()
                        para.style = ref_title_style_name
                        run = para.add_run("参考文献")
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # 添加所有动态生成的引用，根据语言应用不同字体
                        for number, ref_text in unique_references:
                            formatted_ref = f"[{number}] {ref_text}"
                            
                            # 检测引用是否包含中文字符
                            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in ref_text)
                            
                            # 添加参考文献段落
                            add_reference_paragraph(doc, formatted_ref, has_chinese, config)
                    # 更新参考文献数量为合并后的数量
                    references_count = len(unique_references)
            except Exception as e:
                pass
            
            # 确保在保存前文件已经被删除
            if os.path.exists(output_filename):
                try:
                    # 尝试多次删除，因为Word可能暂时锁定文件
                    for attempt in range(10):
                        try:
                            os.remove(output_filename)
                            break
                        except PermissionError:
                            if attempt < 9:
                                print(f"文件被锁定，正在重试删除 ({attempt+1}/10)...")
                                time.sleep(2)  # 等待2秒后重试
                                continue
                            else:
                                print(f"文件锁定失败，请关闭Word后重试")
                                raise
                        except Exception as e:
                            print(f"删除文件时发生错误: {e}")
                            raise
                except Exception as e:
                    print(f"删除现有文件失败: {e}")
                    raise
            # 尝试保存文件
            try:
                doc.save(output_filename)
                # 统计上下标数量
                subscript_count, superscript_count = count_subscript_superscript_in_doc(doc)
                
                print(f"文档生成成功: {output_filename}")
                print(f"- 图片数量: {small_image_count} 张 (大图: {large_image_count} 组)")
                print(f"- 表格数量: {table_counter - 1} 个")
                print(f"- 参考文献数量: {references_count} 篇")
                print(f"- 下标数量: {subscript_count} 个")
                print(f"- 上标数量: {superscript_count} 个")
                
                if duplicate_references:
                    print(f"\n参考文献重复性检测结果:")
                    print(f"- 发现 {len(duplicate_references)} 处重复参考文献")
                    for number, ref_text in duplicate_references:
                        print(f"  重复项 [{number}]: {ref_text}")
                else:
                    print(f"\n参考文献重复性检测结果:")
                    print(f"- 未发现重复参考文献")
                
                if unique_references:
                    print(f"\n最终参考文献列表:")
                    for number, ref_text in unique_references:
                        print(f"  [{number}] {ref_text}")
                
                pdf_output = os.path.splitext(output_filename)[0] + '.pdf'
                convert_word_to_pdf(output_filename, pdf_output)
                    
            except Exception as e:
                print(f"保存文件失败: {e}")
                raise
    except Exception as e:
        import traceback

        traceback.print_exc()